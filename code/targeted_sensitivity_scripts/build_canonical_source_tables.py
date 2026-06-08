from __future__ import annotations

import re
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "Publication" / "paper" / "revision_figures" / "canonical_source_tables"
OUT.mkdir(parents=True, exist_ok=True)

SUPP_TABLES = ROOT / "Supplementary Tables.docx"
DATA_XLSX = ROOT / "Publication" / "paper" / "source_tables" / "scRNA-seq-Benchmark-Data-Overview-20260515.xlsx"


METHOD_FAMILY_MAP = {
    "Linear and Probabilistic Factor Models": "linear/probabilistic",
    "Deep Autoencoders and Generative Models": "deep generative/autoencoder",
    "Graph-Based and Diffusion Geometry Models": "graph/diffusion",
    "Metric Learning and Structure-Aware Embedding": "metric/structure-aware",
}

METHOD_VARIANT_MAP = {
    "TSNE": "t-SNE",
    "ivis": "IVIS",
    "SQuaD_MDS": "SQuaD-MDS",
    "SQuaD_MDS_hybrid": "SQuaD-MDS hybrid",
    "ParametricUMAP50": "Parametric UMAP 50",
    "ParametricUMAP200": "Parametric UMAP 200",
}

VARIANT_PARENT_MAP = {
    "Parametric UMAP 50": "UMAP",
    "Parametric UMAP 200": "UMAP",
    "SQuaD-MDS hybrid": "SQuaD-MDS",
}


def norm_method(value: object) -> str:
    text = str(value)
    return METHOD_VARIANT_MAP.get(text, text)


def parent_method(value: object) -> str:
    text = norm_method(value)
    return VARIANT_PARENT_MAP.get(text, text)


def clean_numeric(value: object) -> float:
    if value is None:
        return np.nan
    text = str(value).strip()
    if not text or text.lower() == "nan":
        return np.nan
    text = text.replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else np.nan


def parse_number_list(value: object) -> list[float]:
    text = str(value)
    nums = re.findall(r"\d+(?:\.\d+)?", text)
    return [float(x) for x in nums]


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT))


def save(df: pd.DataFrame, name: str) -> Path:
    path = OUT / name
    df.to_csv(path, index=False)
    return path


def load_method_manifest() -> pd.DataFrame:
    doc = Document(SUPP_TABLES)
    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:
        vals = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if len(vals) < 8 or not vals[1]:
            continue
        rows.append(
            {
                "method_id": norm_method(vals[1]),
                "parent_method": parent_method(vals[1]),
                "benchmark_scope": "full_26_method_benchmark",
                "method_family_raw": vals[2],
                "method_family": METHOD_FAMILY_MAP.get(vals[2], vals[2]),
                "implementation_principle": vals[3],
                "implementation_language": vals[4],
                "source_or_url": vals[5],
                "publication_year": vals[6],
                "reference": vals[7],
                "is_variant": False,
                "variant_note": "",
            }
        )

    methods = pd.DataFrame(rows)
    known = set(methods["method_id"])

    variant_rows = []
    for variant, parent in VARIANT_PARENT_MAP.items():
        if parent in known:
            base = methods[methods["method_id"] == parent].iloc[0].to_dict()
            base.update(
                {
                    "method_id": variant,
                    "parent_method": parent,
                    "benchmark_scope": "original_result_variant",
                    "is_variant": True,
                    "variant_note": "Shown in original result files as a parameter/workflow variant.",
                }
            )
            variant_rows.append(base)

    scvi = {
        "method_id": "scVI",
        "parent_method": "scVI",
        "benchmark_scope": "targeted_revision_control_only",
        "method_family_raw": "Deep Autoencoders and Generative Models",
        "method_family": "deep generative/autoencoder",
        "implementation_principle": "Deep generative model included only in targeted revision-control analyses.",
        "implementation_language": "Python",
        "source_or_url": "scvi-tools",
        "publication_year": "",
        "reference": "",
        "is_variant": False,
        "variant_note": "Not counted as a full 100-dataset benchmark method in the revised manuscript.",
    }

    out = pd.concat([methods, pd.DataFrame(variant_rows + [scvi])], ignore_index=True, sort=False)
    out["method_order"] = np.arange(1, len(out) + 1)
    return out


def expand_real_s2() -> pd.DataFrame:
    doc = Document(SUPP_TABLES)
    table = doc.tables[1]
    expanded_names = {
        "Nakamura": ["GSE74767/epiblast", "GSE74767/trophectoderm", "GSE74767/ICM", "GSE74767/blastocyst"],
        "Horns": ["GSE100058/DA1_horns", "GSE100058/DC3_VA1d_horns", "GSE100058/horns"],
        "Plass": [
            "GSE103633/combination-1",
            "GSE103633/combination-2",
            "GSE103633/combination-3",
            "GSE103633/epidermis",
            "GSE103633/muscle",
            "GSE103633/neuron",
            "GSE103633/pair-1",
            "GSE103633/pair-2",
            "GSE103633/pair-3",
            "GSE103633/pair-4",
            "GSE103633/parenchyme",
            "GSE103633/phagocyte",
            "GSE103633/pharynx",
        ],
    }
    rows = []
    for idx, row in enumerate(table.rows[1:], start=1):
        vals = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if len(vals) < 8 or not vals[0]:
            continue
        name, source, species, tissue, condition, cell_text, technology, year = vals[:8]
        cells = parse_number_list(cell_text)
        if name in expanded_names:
            names = expanded_names[name]
            for j, dataset_name in enumerate(names):
                rows.append(
                    {
                        "atlas_layer": "manuscript_100_dataset_atlas",
                        "dataset_type": "real",
                        "source_table": "Supplementary Table S2",
                        "source_row": idx,
                        "dataset_label": dataset_name,
                        "dataset_group": name,
                        "source_repository": source,
                        "species": species,
                        "tissue_or_cell_type": tissue,
                        "condition": condition,
                        "cells": cells[j] if j < len(cells) else np.nan,
                        "cell_count_text": cell_text,
                        "sequencing_technology": technology,
                        "year": year,
                        "counting_rule": "expanded_subdataset",
                    }
                )
        else:
            rows.append(
                {
                    "atlas_layer": "manuscript_100_dataset_atlas",
                    "dataset_type": "real",
                    "source_table": "Supplementary Table S2",
                    "source_row": idx,
                    "dataset_label": name,
                    "dataset_group": name,
                    "source_repository": source,
                    "species": species,
                    "tissue_or_cell_type": tissue,
                    "condition": condition,
                    "cells": sum(cells) if name == "Baron" and cells else (cells[0] if cells else np.nan),
                    "cell_count_text": cell_text,
                    "sequencing_technology": technology,
                    "year": year,
                    "counting_rule": "single_row_dataset" if name != "Baron" else "baron_human_mouse_counted_as_one_dataset",
                }
            )
    return pd.DataFrame(rows)


def load_excel_dataset_details() -> tuple[pd.DataFrame, pd.DataFrame]:
    real = pd.read_excel(DATA_XLSX, sheet_name=1)
    real = real[real["category"].isin(["SIMLR", "VASC", "benchmarker", "scDesign3", "TI"]) & real["dataset"].notna()].copy()
    real = real[real["dataset"].astype(str).str.contains("/")].copy()
    real["dataset_type"] = "real"
    real["atlas_layer"] = "explicit_excel_detail_or_repository_snapshot"
    for col in ["cells", "genes", "sparsity_pct", "cell_types", "size_mb"]:
        real[col] = real[col].map(clean_numeric)

    sim = pd.read_excel(DATA_XLSX, sheet_name=2)
    valid_groups = {"default", "batch", "cell", "celltype", "de", "de_prob", "dropout", "gene", "out"}
    sim = sim[sim["param_group"].isin(valid_groups) & sim["cells"].notna() & sim["size_mb"].notna()].copy()
    sim["dataset_type"] = "simulated"
    sim["atlas_layer"] = "manuscript_100_dataset_atlas"
    for col in ["cells", "genes", "sparsity_pct", "n_batches", "n_groups", "size_mb"]:
        sim[col] = sim[col].map(clean_numeric)
    sim["dataset_label"] = sim["param_group"].astype(str) + "_" + sim["param_value"].astype(str)
    return real, sim


def result_availability() -> pd.DataFrame:
    rows = []
    for base_name, root in [
        ("results_real", ROOT / "scRNA-DRComparison-main" / "results" / "data" / "real"),
        ("datasets_real", ROOT / "scRNA-DRComparison-main" / "datasets" / "real"),
        ("results_simulated", ROOT / "scRNA-DRComparison-main" / "results" / "data" / "simulate"),
        ("datasets_simulated", ROOT / "scRNA-DRComparison-main" / "datasets" / "simulate"),
    ]:
        if not root.exists():
            continue
        for meta in sorted(root.rglob("cell_metadata.csv")):
            dataset_id = str(meta.parent.relative_to(root)).replace("\\", "/")
            rows.append(
                {
                    "availability_source": base_name,
                    "dataset_id": dataset_id,
                    "has_cell_metadata": True,
                    "has_counts_matrix": (meta.parent / "counts_matrix.csv").exists(),
                    "relative_path": rel(meta.parent),
                }
            )
    return pd.DataFrame(rows)


def load_revision_subset_manifest() -> pd.DataFrame:
    path = ROOT / "revision_benchmark" / "experiments" / "datasets_manifest.csv"
    df = pd.read_csv(path)
    df["atlas_layer"] = "targeted_revision_control_subset"
    df["dataset_type"] = df["dataset_group"].map({"synthetic": "simulated", "downsampling": "downsampling"}).fillna(df["dataset_group"])
    for col in ["n_cells", "n_genes"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


def load_score_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    rows = []
    score_files = {
        "local": ROOT / "metric" / "score" / "local" / "local.csv",
        "global": ROOT / "metric" / "score" / "global" / "global.csv",
        "kmeans": ROOT / "metric" / "score" / "cluster" / "kmeans.csv",
        "louvain": ROOT / "metric" / "score" / "cluster" / "louvain.csv",
        "spectral": ROOT / "metric" / "score" / "cluster" / "spectral.csv",
        "runtime_score": ROOT / "metric" / "score" / "efficiency" / "time.csv",
        "memory_score": ROOT / "metric" / "score" / "efficiency" / "memory.csv",
    }
    for metric, path in score_files.items():
        df = pd.read_csv(path)
        df["method_raw"] = df["Method"].astype(str)
        df["method_id"] = df["method_raw"].map(norm_method)
        df["parent_method"] = df["method_id"].map(parent_method)
        df["score_domain"] = metric
        df["source_file"] = rel(path)
        rows.append(df[["method_raw", "method_id", "parent_method", "score_domain", "score", "source_file"]])
    score_long = pd.concat(rows, ignore_index=True)

    stability_rows = []
    for path in sorted((ROOT / "metric" / "score" / "stability").glob("*.csv")):
        df = pd.read_csv(path)
        df["method_raw"] = df["Method"].astype(str)
        df["method_id"] = df["method_raw"].map(norm_method)
        df["parent_method"] = df["method_id"].map(parent_method)
        df["perturbation_axis"] = path.stem
        df["source_file"] = rel(path)
        stability_rows.append(df[["method_raw", "method_id", "parent_method", "perturbation_axis", "score", "source_file"]])
    stability = pd.concat(stability_rows, ignore_index=True)
    stability_median = stability.groupby(["method_id", "parent_method"], as_index=False)["score"].median()
    stability_median["method_raw"] = stability_median["method_id"]
    stability_median["score_domain"] = "stability_median"
    stability_median["source_file"] = "metric/score/stability/*.csv"
    score_long = pd.concat(
        [score_long, stability_median[["method_raw", "method_id", "parent_method", "score_domain", "score", "source_file"]]],
        ignore_index=True,
    )
    score_matrix = score_long.pivot_table(index="method_id", columns="score_domain", values="score", aggfunc="mean")
    preferred = ["local", "global", "kmeans", "louvain", "spectral", "runtime_score", "memory_score", "stability_median"]
    score_matrix = score_matrix.reindex(columns=[c for c in preferred if c in score_matrix.columns])
    score_matrix["overall_mean"] = score_matrix.mean(axis=1, skipna=True)
    score_matrix = score_matrix.sort_values("overall_mean", ascending=False).reset_index()
    return score_long, score_matrix, stability


def parse_topology_raw() -> pd.DataFrame:
    rows = []
    for path in sorted((ROOT / "metric" / "topo").rglob("dr*.csv")):
        rel_parts = path.relative_to(ROOT / "metric" / "topo").parts
        if len(rel_parts) < 3:
            continue
        category = rel_parts[0]
        dataset_id = "/".join(rel_parts[1:-1])
        table_id = path.stem
        df = pd.read_csv(path)
        if "Method" not in df.columns:
            continue
        df["method_raw"] = df["Method"].astype(str)
        df["method_id"] = df["method_raw"].map(norm_method)
        df["parent_method"] = df["method_id"].map(parent_method)
        value_cols = [c for c in df.columns if c not in {"Method", "method_raw", "method_id", "parent_method"}]
        long = df.melt(
            id_vars=["method_raw", "method_id", "parent_method"],
            value_vars=value_cols,
            var_name="metric",
            value_name="value",
        )
        long["dataset_category"] = category
        long["dataset_id"] = dataset_id
        long["topology_table"] = table_id
        long["source_file"] = rel(path)
        rows.append(long)
    return pd.concat(rows, ignore_index=True) if rows else pd.DataFrame()


def parse_clustering_raw() -> pd.DataFrame:
    rows = []
    for path in sorted((ROOT / "metric" / "cluster").rglob("indicators/*.csv")):
        rel_parts = path.relative_to(ROOT / "metric" / "cluster").parts
        if len(rel_parts) < 4:
            continue
        category = rel_parts[0]
        dataset_id = "/".join(rel_parts[1:-2])
        name = path.stem
        if "_" not in name:
            continue
        algorithm, metric = name.split("_", 1)
        df = pd.read_csv(path)
        if "Method" not in df.columns or metric not in df.columns:
            continue
        out = df[["Method", metric]].copy()
        out = out.rename(columns={"Method": "method_raw", metric: "value"})
        out["method_id"] = out["method_raw"].map(norm_method)
        out["parent_method"] = out["method_id"].map(parent_method)
        out["dataset_category"] = category
        out["dataset_id"] = dataset_id
        out["clustering_algorithm"] = algorithm
        out["metric"] = metric
        out["source_file"] = rel(path)
        rows.append(out)
    return pd.concat(rows, ignore_index=True) if rows else pd.DataFrame()


def load_efficiency() -> pd.DataFrame:
    rows = []
    for path in sorted((ROOT / "metric" / "efficiency").glob("cell_*.csv")):
        df = pd.read_csv(path)
        n_cells = int(path.stem.split("_")[1])
        df["method_raw"] = df["Method"].astype(str)
        df["method_id"] = df["method_raw"].map(norm_method)
        df["parent_method"] = df["method_id"].map(parent_method)
        df["n_cells"] = n_cells
        df["peak_memory_gb"] = pd.to_numeric(df["PeakMemory(gb)"], errors="coerce")
        df["runtime_seconds"] = pd.to_numeric(df["Time(s)"], errors="coerce")
        df["source_file"] = rel(path)
        rows.append(df[["method_raw", "method_id", "parent_method", "n_cells", "peak_memory_gb", "runtime_seconds", "source_file"]])
    return pd.concat(rows, ignore_index=True)


def load_revision_controls() -> dict[str, pd.DataFrame]:
    paths = {
        "revision_scvi_control": ROOT / "revision_benchmark" / "results" / "metrics" / "WP1_scVI_local_scVI_metrics.csv",
        "revision_dimension_sensitivity": ROOT / "revision_benchmark" / "results" / "source_data" / "wp2_dimension_sensitivity_long.csv",
        "revision_visualization_workflow": ROOT / "revision_benchmark" / "results" / "source_data" / "wp3_visualization_workflow_long.csv",
        "revision_input_gene_sensitivity": ROOT / "revision_benchmark" / "results" / "source_data" / "wp4_input_gene_sensitivity_long.csv",
    }
    out = {}
    for name, path in paths.items():
        df = pd.read_csv(path)
        if "method" in df.columns:
            df["method_raw"] = df["method"].astype(str)
            df["method_id"] = df["method_raw"].map(norm_method)
            df["parent_method"] = df["method_id"].map(parent_method)
        df["control_layer"] = name
        df["source_file"] = rel(path)
        out[name] = df
    return out


def build_summary() -> pd.DataFrame:
    rows = []
    for path in sorted(OUT.glob("*.csv")):
        df = pd.read_csv(path)
        rows.append({"table": path.name, "rows": len(df), "columns": len(df.columns), "path": rel(path)})
    return pd.DataFrame(rows)


def main() -> None:
    method_manifest = load_method_manifest()
    real_s2 = expand_real_s2()
    real_xlsx, sim_xlsx = load_excel_dataset_details()
    availability = result_availability()
    revision_manifest = load_revision_subset_manifest()
    score_long, score_matrix, stability = load_score_tables()
    topology_raw = parse_topology_raw()
    clustering_raw = parse_clustering_raw()
    efficiency = load_efficiency()
    revision_controls = load_revision_controls()

    save(method_manifest, "canonical_method_manifest.csv")
    save(real_s2, "canonical_real_dataset_atlas_from_supplementary_table_s2.csv")
    save(real_xlsx, "excel_real_dataset_detail_rows.csv")
    save(sim_xlsx, "canonical_simulated_dataset_atlas_from_excel.csv")
    save(availability, "repository_dataset_availability.csv")
    save(revision_manifest, "revision_control_dataset_manifest.csv")
    save(score_long, "original_score_long.csv")
    save(score_matrix, "original_score_matrix.csv")
    save(stability, "original_stability_score_long.csv")
    save(topology_raw, "original_topology_raw_long.csv")
    save(clustering_raw, "original_clustering_raw_long.csv")
    save(efficiency, "original_efficiency_scaling_long.csv")

    for name, df in revision_controls.items():
        save(df, f"{name}.csv")

    summary = build_summary()
    save(summary, "canonical_source_table_manifest.csv")

    checks = pd.DataFrame(
        [
            {"check": "full_benchmark_methods", "value": int((method_manifest["benchmark_scope"] == "full_26_method_benchmark").sum()), "expected": 26},
            {"check": "real_manuscript_atlas_rows", "value": len(real_s2), "expected": 50},
            {"check": "real_excel_explicit_rows", "value": len(real_xlsx), "expected": 49},
            {"check": "simulated_atlas_rows", "value": len(sim_xlsx), "expected": 50},
            {"check": "revision_manifest_rows", "value": len(revision_manifest), "expected": 60},
            {"check": "topology_raw_rows", "value": len(topology_raw), "expected": "nonzero"},
            {"check": "clustering_raw_rows", "value": len(clustering_raw), "expected": "nonzero"},
            {"check": "efficiency_rows", "value": len(efficiency), "expected": "nonzero"},
        ]
    )
    save(checks, "canonical_source_table_checks.csv")

    print(f"Wrote canonical source tables to {OUT}")
    print(checks.to_string(index=False))


if __name__ == "__main__":
    main()
