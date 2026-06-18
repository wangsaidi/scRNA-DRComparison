from __future__ import annotations

from pathlib import Path
import re
import shutil

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


def find_project_root(start: Path) -> Path:
    for parent in [start.resolve(), *start.resolve().parents]:
        if (parent / "Publication" / "paper").exists() and (parent / "metadata").exists():
            return parent
    raise RuntimeError(f"Could not locate project root from {start}")


ROOT = find_project_root(Path(__file__))
OUT = ROOT / "Publication/paper/revision_tables/final_supplementary_tables_package_20260606"
TABLE_OUT = OUT / "tables"
CSV_OUT = OUT / "source_tables"
QA_OUT = OUT / "qa"
CODE_OUT = OUT / "code"
for path in [TABLE_OUT, CSV_OUT, QA_OUT, CODE_OUT]:
    path.mkdir(parents=True, exist_ok=True)

REDESIGNED_SRC = ROOT / "Publication/paper/revision_figures/redesigned_python_figure_package/source_data"
MAIN_SRC = ROOT / "Publication/paper/revision_figures/main_figures_final_package_20260605/source_data"
ATLAS_PACKAGE = ROOT / "Publication/paper/revision_figures/final_supplementary_atlas_package_20260606"
ATLAS_PANEL = ATLAS_PACKAGE / "source_data/panel_data"
ATLAS_QA = ATLAS_PACKAGE / "qa"
SUPP_UPDATE = ROOT / "Publication/paper/revision_figures/supplementary_update/source_tables"


CANONICAL_METHODS = [
    "PCA",
    "GLMPCA",
    "pCMF",
    "ZIFA",
    "scGBM",
    "SSNMDI",
    "tGPLVM",
    "VAE",
    "scvis",
    "VASC",
    "SAUCIE",
    "scScope",
    "SCDRHA",
    "scGAE",
    "DREAM",
    "DRA",
    "UMAP",
    "SIMLR",
    "PHATE",
    "EDGE",
    "SPDR",
    "IVIS",
    "PaCMAP",
    "t-SNE",
    "TriMap",
    "SQuaD-MDS",
]
ORDER_MAP = {method: i + 1 for i, method in enumerate(CANONICAL_METHODS)}
SOURCE_COLLECTION_LABELS = {
    "SIMLR": "SIMLR reference datasets",
    "VASC": "VASC reference datasets",
    "benchmarker": "real-data benchmark collection",
    "scDesign3": "scDesign3 reference datasets",
    "TI": "trajectory-inference reference datasets",
    "simulate": "simulated perturbation atlas",
}


def supplementary_label(table_id: str) -> str:
    return f"Supplementary Table {table_id.split('_', 1)[0]}"


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def normalize_method(series: pd.Series) -> pd.Series:
    return series.astype(str).replace(
        {
            "TSNE": "t-SNE",
            "SQuaD_MDS": "SQuaD-MDS",
            "ivis": "IVIS",
            "ParametricUMAP50": "Parametric UMAP 50",
            "ParametricUMAP200": "Parametric UMAP 200",
        }
    )


def add_method_order(df: pd.DataFrame, col: str = "method_id") -> pd.DataFrame:
    out = df.copy()
    if col in out.columns:
        out[col] = normalize_method(out[col])
        out["method_order"] = out[col].map(ORDER_MAP)
        out = out.sort_values(["method_order", col], na_position="last")
    return out


def public_text(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    replacements = {
        "full_26_method_benchmark": "Full benchmark",
        "original_result_variant": "Result variant",
        "targeted_revision_control_only": "Targeted reference",
        "WP1_scVI_local": "scVI reference",
        "WP2_dimension_sensitivity": "latent-dimension sensitivity",
        "WP3_visualization_workflow": "visualization-workflow comparison",
        "WP4_input_gene_sensitivity": "input-gene sensitivity",
        "scVI_reference_analysis": "scVI reference analysis",
        "dimension_sensitivity": "latent-dimension sensitivity",
        "visualization_workflow": "visualization-workflow comparison",
        "input_gene_sensitivity": "input-gene sensitivity",
    }
    for col in out.select_dtypes(include="object").columns:
        s = out[col].astype("string")
        for old, new in replacements.items():
            s = s.str.replace(old, new, regex=False)
        out[col] = s
    return out


def table_index_rows(tables: dict[str, tuple[str, pd.DataFrame, str, str]]) -> pd.DataFrame:
    rows = []
    for table_id, (title, df, primary_figures, purpose) in tables.items():
        rows.append(
            {
                "supplementary_table": supplementary_label(table_id),
                "table_id": table_id,
                "title": title,
                "rows": len(df),
                "columns": len(df.columns),
                "primary_figures_or_sections": primary_figures,
                "purpose": purpose,
            }
        )
    return pd.DataFrame(rows)


def make_method_catalog() -> pd.DataFrame:
    df = read_csv(REDESIGNED_SRC / "Supplementary_Figure_S1_method_catalog_source_data.csv")
    df = public_text(df)
    df["parent_method"] = normalize_method(df["parent_method"])
    df["method_id"] = normalize_method(df["method_id"])
    df["method_order"] = df["parent_method"].map(ORDER_MAP).fillna(df["method_order"])
    scvi = df["method_id"].eq("scVI")
    df.loc[scvi, "publication_year"] = 2018
    df.loc[scvi, "reference"] = "(Lopez et al., 2018; doi:10.1038/s41592-018-0229-2)"
    df["variant_note"] = df["variant_note"].fillna("not applicable").replace("", "not applicable")
    columns = [
        "method_order",
        "method_id",
        "parent_method",
        "benchmark_scope",
        "method_family",
        "implementation_principle",
        "implementation_language",
        "source_or_url",
        "publication_year",
        "reference",
        "is_variant",
        "variant_note",
    ]
    return df[[c for c in columns if c in df.columns]].sort_values(["method_order", "method_id"], na_position="last")


def make_dataset_100() -> pd.DataFrame:
    df = read_csv(REDESIGNED_SRC / "Supplementary_Figure_S2_full_dataset_atlas_source_data.csv")
    columns = [
        "atlas_layer",
        "dataset_type",
        "source_table",
        "source_row",
        "dataset_label",
        "dataset_group",
        "source_repository",
        "species",
        "tissue_or_cell_type",
        "condition",
        "cells",
        "cell_count_text",
        "sequencing_technology",
        "year",
        "counting_rule",
        "source",
        "param_group",
        "param_value",
        "genes",
        "sparsity_pct",
        "n_batches",
        "n_groups",
        "size_mb",
    ]
    return public_text(df[[c for c in columns if c in df.columns]])


def make_real_detail() -> pd.DataFrame:
    atlas = make_dataset_100()
    atlas = atlas[atlas["dataset_type"].str.lower().eq("real")].copy()
    detail = read_csv(REDESIGNED_SRC / "Supplementary_Figure_S3_real_dataset_landscape_source_data.csv")
    detail = detail.drop(columns=[c for c in detail.columns if c.startswith("Unnamed")], errors="ignore")

    def key(value: object) -> str:
        tail = str(value).split("/")[-1].lower()
        return re.sub(r"[^a-z0-9]+", "", tail)

    detail = detail.copy()
    detail["_match_key"] = detail["dataset"].map(key)
    detail = detail.drop_duplicates("_match_key", keep="first")

    atlas["_match_key"] = atlas["dataset_label"].map(key)
    merged = atlas.merge(
        detail.add_prefix("detail_"),
        left_on="_match_key",
        right_on="detail__match_key",
        how="left",
    )
    merged["metadata_status"] = np.where(
        merged["detail_dataset"].notna(),
        "Detailed metadata available",
        "Atlas-level metadata only",
    )
    merged["metadata_note"] = ""
    merged.loc[
        merged["dataset_label"].eq("Baron"),
        "metadata_note",
    ] = "Baron human and mouse entries are counted as one manuscript real dataset in the 100-dataset atlas."
    merged.loc[
        merged["metadata_status"].eq("Atlas-level metadata only"),
        "metadata_note",
    ] = "Dataset is retained in the 50-real-dataset manuscript atlas; fields not present in the detailed snapshot should be reported as not available rather than inferred."
    merged["metadata_note"] = merged["metadata_note"].replace("", "No additional note.")

    out = pd.DataFrame(
        {
            "source_row": merged["source_row"],
            "dataset_label": merged["dataset_label"],
            "dataset_group": merged["dataset_group"],
            "source_repository": merged["source_repository"],
            "species": merged["species"],
            "tissue_or_cell_type": merged["tissue_or_cell_type"],
            "condition": merged["condition"],
            "cells_atlas": merged["cells"],
            "cell_count_text": merged["cell_count_text"],
            "sequencing_technology": merged["sequencing_technology"],
            "year": merged["year"],
            "counting_rule": merged["counting_rule"],
            "detail_source_dataset": merged["detail_dataset"],
            "genes_detail": merged["detail_genes"],
            "sparsity_pct_detail": merged["detail_sparsity_pct"],
            "cell_types_detail": merged["detail_cell_types"],
            "size_mb_detail": merged["detail_size_mb"],
            "metadata_status": merged["metadata_status"],
            "metadata_note": merged["metadata_note"],
        }
    )
    out["metadata_note"] = out["metadata_note"].fillna("No additional note.").replace("", "No additional note.")
    return public_text(out.sort_values("source_row", na_position="last"))


def make_sim_detail() -> pd.DataFrame:
    df = read_csv(REDESIGNED_SRC / "Supplementary_Figure_S4_simulated_parameter_landscape_source_data.csv")
    return public_text(df)


def make_metric_inventory() -> pd.DataFrame:
    rows = []
    score_path = SUPP_UPDATE / "Supplementary_Table_S5_Figure3_score_component_inventory.csv"
    if score_path.exists():
        score = read_csv(score_path)
        for _, row in score.iterrows():
            rows.append(
                {
                    "inventory_layer": "profile score component",
                    "domain_or_algorithm": row.get("score_domain"),
                    "metric_component": row.get("metric_component"),
                    "display_definition": row.get("display_definition"),
                    "raw_direction": row.get("raw_direction"),
                    "component_order": row.get("component_order"),
                    "profile_weight_per_component": row.get("profile_weight_per_component"),
                    "source": "Figure 3 score construction",
                }
            )

    struct = read_csv(REDESIGNED_SRC / "Figure_4_structure_preservation_source_data.csv")
    struct = struct[struct["metric"].notna()].copy()
    for metric in sorted(struct["metric"].dropna().unique()):
        domain = "local/topology" if any(metric.startswith(x) for x in ["nkr", "aji", "T_", "C_", "Mrre", "nh", "knn", "svm"]) else "global/geometry"
        direction = "cost or rank-error" if metric.startswith("Mrre") else "benefit"
        rows.append(
            {
                "inventory_layer": "structure preservation metric",
                "domain_or_algorithm": domain,
                "metric_component": metric,
                "display_definition": metric,
                "raw_direction": direction,
                "component_order": np.nan,
                "profile_weight_per_component": np.nan,
                "source": "Figure 4 and Supplementary Figures S6-S8",
            }
        )

    clust = read_csv(REDESIGNED_SRC / "Figure_5_clustering_concordance_source_data.csv")
    for algorithm in sorted(clust["clustering_algorithm"].dropna().unique()):
        for metric in sorted(clust["metric"].dropna().unique()):
            rows.append(
                {
                    "inventory_layer": "clustering concordance metric",
                    "domain_or_algorithm": algorithm,
                    "metric_component": metric,
                    "display_definition": f"{metric} under {algorithm} clustering",
                    "raw_direction": "benefit",
                    "component_order": np.nan,
                    "profile_weight_per_component": np.nan,
                    "source": "Figure 5 and Supplementary Figure S9",
                }
            )
    out = pd.DataFrame(rows).drop_duplicates()
    for col in ["component_order", "profile_weight_per_component"]:
        if col in out.columns:
            out[col] = out[col].where(out[col].notna(), "not applicable")
    return public_text(out)


def make_score_matrix() -> pd.DataFrame:
    df = read_csv(MAIN_SRC / "figure3/Figure_3_completed_score_matrix.csv")
    df = add_method_order(df, "method_id")
    ordered = ["method_order", "method_id", "family"] + [c for c in df.columns if c not in {"method_order", "method_id", "family"}]
    return df[ordered]


def make_structure_summary() -> pd.DataFrame:
    df = read_csv(REDESIGNED_SRC / "Figure_4_structure_preservation_source_data.csv")
    df = df[df["metric"].notna() & df["value"].notna()].copy()
    df["parent_method"] = normalize_method(df["parent_method"])
    summary = (
        df.groupby(["dataset_category", "metric"], observed=True)
        .agg(
            datasets=("dataset_id", "nunique"),
            methods=("parent_method", "nunique"),
            records=("value", "size"),
            median_value=("value", "median"),
            q25_value=("value", lambda x: x.quantile(0.25)),
            q75_value=("value", lambda x: x.quantile(0.75)),
        )
        .reset_index()
        .sort_values(["dataset_category", "metric"])
    )
    summary["dataset_category"] = summary["dataset_category"].replace(SOURCE_COLLECTION_LABELS)
    summary = summary.rename(columns={"dataset_category": "source_collection"})
    return public_text(summary)


def make_clustering_summary() -> pd.DataFrame:
    df = read_csv(REDESIGNED_SRC / "Figure_5_clustering_concordance_source_data.csv")
    df["parent_method"] = normalize_method(df["parent_method"])
    summary = (
        df.dropna(subset=["value"])
        .groupby(["dataset_category", "clustering_algorithm", "metric"], observed=True)
        .agg(
            datasets=("dataset_id", "nunique"),
            methods=("parent_method", "nunique"),
            records=("value", "size"),
            median_value=("value", "median"),
            q25_value=("value", lambda x: x.quantile(0.25)),
            q75_value=("value", lambda x: x.quantile(0.75)),
        )
        .reset_index()
        .sort_values(["dataset_category", "clustering_algorithm", "metric"])
    )
    summary["dataset_category"] = summary["dataset_category"].replace(SOURCE_COLLECTION_LABELS)
    summary = summary.rename(columns={"dataset_category": "source_collection"})
    return public_text(summary)


def make_revision_experiment_manifest() -> pd.DataFrame:
    full = read_csv(MAIN_SRC / "figure6/Figure_6_targeted_sensitivity_controls_full_source_data.csv")
    full = public_text(full)
    rows = []
    for layer, group in full.groupby("control_layer", dropna=False, observed=True):
        layer_label = str(layer).replace(" comparison_comparison", " comparison")
        parameter_parts = []
        if group["dimension"].notna().any():
            parameter_parts.append("latent dimensions: " + "; ".join(map(str, sorted(group["dimension"].dropna().astype(int).unique()))))
        if "workflow" in group.columns and group["workflow"].notna().any():
            parameter_parts.append("workflows: " + "; ".join(sorted(group["workflow"].dropna().astype(str).unique())))
        if "hvg_requested" in group.columns and group["hvg_requested"].notna().any():
            parameter_parts.append("requested HVGs: " + "; ".join(map(str, sorted(group["hvg_requested"].dropna().astype(int).unique()))))
        rows.append(
            {
                "record_type": "experiment manifest",
                "item": layer_label,
                "value": f"{len(group)} source records; {group['parent_method'].dropna().nunique()} methods; {group['dataset_id'].dropna().nunique()} datasets",
                "method_scope": "; ".join(sorted(group["parent_method"].dropna().astype(str).unique())),
                "dataset_scope": "; ".join(sorted(group["dataset_id"].dropna().astype(str).unique())[:15]),
                "parameter_scope": "; ".join(parameter_parts) if parameter_parts else "not applicable",
                "primary_figure": "Figure 6; Supplementary Figures S11-S14",
                "evidence_or_note": "Defines the targeted sensitivity experiment scope used to evaluate scVI, latent dimension, visualization workflow, and input-gene sensitivity.",
            }
        )
    qa = read_csv(MAIN_SRC / "figure6/Figure_6_targeted_sensitivity_controls_polished_qa_summary.csv")
    qa_long = qa.T.reset_index()
    qa_long.columns = ["item", "value"]
    qa_rows = []
    for _, row in qa_long.iterrows():
        qa_rows.append(
            {
                "record_type": "Figure 6 QA summary",
                "item": str(row["item"]).replace("control_response", "control_effect"),
                "value": row["value"],
                "method_scope": "not applicable",
                "dataset_scope": "not applicable",
                "parameter_scope": "not applicable",
                "primary_figure": "Figure 6",
                "evidence_or_note": "Computed from the polished Figure 6 source-data QA summary.",
            }
        )
    return public_text(pd.DataFrame(rows + qa_rows))


def make_robustness_summary() -> pd.DataFrame:
    score = make_score_matrix()
    axes = [
        "cell_number",
        "gene_number",
        "celltype_number",
        "dropout",
        "batch_number",
        "batch_strength",
        "de_prob",
        "de_strength",
        "out",
        "stability_median",
    ]
    columns = ["method_order", "method_id", "family"] + [c for c in axes if c in score.columns]
    return score[columns]


def make_scalability_audit() -> pd.DataFrame:
    completion = read_csv(MAIN_SRC / "figure8/Figure_8_completion_missingness_audit_source_data.csv")
    install = read_csv(MAIN_SRC / "figure8/Figure_8_install_manifest_full26_source_data.csv")
    efficiency = read_csv(MAIN_SRC / "figure8/Figure_8_efficiency_scaling_full26_source_data.csv")
    install = install.rename(columns={"method": "method_id_from_install"})
    install["method_id"] = install["method_id"].fillna(install["method_id_from_install"])
    install["method_id"] = normalize_method(install["method_id"])
    completion["method_id"] = normalize_method(completion["method_id"])
    efficiency["parent_method"] = normalize_method(efficiency["parent_method"])
    eff_summary = (
        efficiency.groupby("parent_method", observed=True)
        .agg(
            efficiency_rows=("runtime_seconds", "size"),
            median_runtime_seconds=("runtime_seconds", "median"),
            max_runtime_seconds=("runtime_seconds", "max"),
            median_peak_memory_gb=("peak_memory_gb", "median"),
            max_peak_memory_gb=("peak_memory_gb", "max"),
            largest_efficiency_scale_cells=("n_cells", "max"),
        )
        .reset_index()
        .rename(columns={"parent_method": "method_id"})
    )
    out = completion.merge(install[["method_id", "language", "role", "install_channel", "package_or_repo", "environment", "status", "notes"]], on="method_id", how="left")
    out = out.merge(eff_summary, on="method_id", how="left")
    if "missing_cell_levels" in out.columns:
        out["missing_cell_levels"] = out["missing_cell_levels"].fillna("none")
    out = out.rename(columns={col: f"completed_{col}_cells" for col in out.columns if str(col).isdigit()})
    out = add_method_order(public_text(out), "method_id")
    ordered = ["method_order", "method_id"] + [c for c in out.columns if c not in {"method_order", "method_id"}]
    return out[ordered]


def make_supp_figure_panel_map() -> pd.DataFrame:
    return public_text(read_csv(ATLAS_QA / "panel_role_audit.csv"))


def make_supplementary_evidence_coverage() -> pd.DataFrame:
    source = public_text(read_csv(ATLAS_QA / "old_supplement_to_new_atlas_mapping.csv"))
    source = source.rename(
        columns={
            "new_figure": "supplementary_figure",
            "consolidated_role": "evidence_layer",
        }
    )
    coverage = (
        source.groupby(["supplementary_figure", "evidence_layer"], dropna=False)
        .size()
        .reset_index(name="evidence_units")
        .sort_values(["supplementary_figure", "evidence_layer"])
    )
    coverage["coverage_scope"] = coverage["evidence_layer"].map(
        {
            "local structure: NKR/AJI/MRRE": "local-neighborhood preservation metrics",
            "trustworthiness/continuity/NH/KNN/SVM": "local-manifold and neighborhood-classification metrics",
            "global/class geometry": "global geometry and class-separation metrics",
            "clustering concordance": "cluster-label concordance metrics",
            "synthetic perturbation": "controlled simulated perturbation axes",
        }
    )
    coverage["source_layer_type"] = "metric-level supplementary evidence"
    coverage["reader_use"] = (
        "Use this row to locate the supplementary figure that summarizes this evidence layer."
    )
    return coverage[
        [
            "supplementary_figure",
            "evidence_layer",
            "coverage_scope",
            "evidence_units",
            "source_layer_type",
            "reader_use",
        ]
    ]


def make_source_file_manifest() -> pd.DataFrame:
    rows = []
    search_dirs = [
        ("supplementary_figure_panel_data", ATLAS_PANEL),
        ("main_figure_source_data", MAIN_SRC),
        ("compiled_figure_source_data", REDESIGNED_SRC),
        ("supplementary_table_source_data", SUPP_UPDATE),
    ]
    for role, directory in search_dirs:
        if not directory.exists():
            continue
        for path in sorted(directory.rglob("*.csv")):
            try:
                df_head = pd.read_csv(path, nrows=5)
                row_count = sum(1 for _ in path.open("rb")) - 1
                col_count = len(df_head.columns)
                columns_preview = "; ".join(map(str, df_head.columns[:12]))
            except Exception as exc:
                row_count = np.nan
                col_count = np.nan
                columns_preview = f"unreadable: {exc}"
            display_path = str(path.relative_to(ROOT))
            display_path = display_path.replace(
                "Publication\\paper\\revision_figures\\main_figures_final_package_20260605\\",
                "main_figures_final_package\\",
            )
            display_path = display_path.replace(
                "Publication\\paper\\revision_figures\\final_supplementary_atlas_package_20260606\\",
                "final_supplementary_atlas_package\\",
            )
            display_path = display_path.replace(
                "Publication\\paper\\revision_figures\\redesigned_python_figure_package\\",
                "compiled_figure_source_package\\",
            )
            display_path = display_path.replace(
                "Publication\\paper\\revision_figures\\supplementary_update\\",
                "supplementary_table_source_archive\\",
            )
            display_path = display_path.replace("revision_control_dataset_manifest", "targeted_control_dataset_manifest")
            display_path = display_path.replace("revision_dimension_sensitivity", "targeted_dimension_sensitivity")
            display_path = display_path.replace("revision_input_gene_sensitivity", "targeted_input_gene_sensitivity")
            display_path = display_path.replace("revision_scvi_control", "targeted_scvi_reference")
            display_path = display_path.replace("revision_visualization_workflow", "targeted_visualization_workflow")
            rows.append(
                {
                    "source_role": role,
                    "relative_path": display_path,
                    "rows": row_count,
                    "columns": col_count,
                    "columns_preview": columns_preview,
                }
            )
    return pd.DataFrame(rows)


def make_data_route_audit(source_manifest: pd.DataFrame) -> pd.DataFrame:
    rows = [
        {
            "data_or_output_class": "Method catalogue and implementation metadata",
            "access_route": "within paper or supplement",
            "location_in_package": "Supplementary Table S1; source_tables/S1_method_catalog.csv",
            "risk_or_action": "No action required.",
        },
        {
            "data_or_output_class": "100-dataset benchmark metadata",
            "access_route": "within paper or supplement; reused public source metadata",
            "location_in_package": "Supplementary Tables S2-S4; source tables",
            "risk_or_action": "Confirm final public repository/accession wording for reused datasets in Data Availability.",
        },
        {
            "data_or_output_class": "Benchmark metric outputs and score matrices",
            "access_route": "within paper or supplement",
            "location_in_package": "Supplementary Tables S5-S10; source_tables and final supplementary figure package",
            "risk_or_action": "Large raw per-dataset metrics are summarized in tables and retained as source data.",
        },
        {
            "data_or_output_class": "Targeted sensitivity experiments",
            "access_route": "within paper or supplement",
            "location_in_package": "Supplementary Table S11; Figure 6 source data",
            "risk_or_action": "scVI is marked as targeted reference analysis, not part of the 26-method full benchmark.",
        },
        {
            "data_or_output_class": "Code and figure/table reproduction scripts",
            "access_route": "within supplement/package; repository to be confirmed",
            "location_in_package": "code/ and package README",
            "risk_or_action": "Confirm final code repository or archive DOI before submission.",
        },
    ]
    rows.append(
        {
            "data_or_output_class": "Source file manifest",
            "access_route": "within supplement/package",
            "location_in_package": f"Supplementary Table S15; {len(source_manifest)} source files indexed",
            "risk_or_action": "No action required.",
        }
    )
    return pd.DataFrame(rows)


def write_csvs(tables: dict[str, tuple[str, pd.DataFrame, str, str]]) -> None:
    for old_csv in CSV_OUT.glob("S*.csv"):
        old_csv.unlink()
    for table_id, (_, df, _, _) in tables.items():
        df.to_csv(CSV_OUT / f"{table_id}.csv", index=False, encoding="utf-8-sig")


def write_excel(tables: dict[str, tuple[str, pd.DataFrame, str, str]]) -> Path:
    xlsx = TABLE_OUT / "Supplementary_Tables_final_20260606.xlsx"
    with pd.ExcelWriter(xlsx, engine="openpyxl") as writer:
        index = table_index_rows(tables)
        index.to_excel(writer, sheet_name="Table_Index", index=False)
        for table_id, (title, df, _, _) in tables.items():
            sheet = table_id[:31]
            df.to_excel(writer, sheet_name=sheet, index=False)
    wb = load_workbook(xlsx)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for col_idx, column_cells in enumerate(ws.columns, start=1):
            max_len = 0
            for cell in column_cells[:200]:
                value = "" if cell.value is None else str(cell.value)
                max_len = max(max_len, min(len(value), 60))
            ws.column_dimensions[get_column_letter(col_idx)].width = max(10, min(max_len + 2, 42))
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(xlsx)
    return xlsx


def add_small_table(doc: Document, df: pd.DataFrame, max_rows: int = 12) -> None:
    preview = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(preview.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(preview.columns):
        hdr[idx].text = str(col)
    for _, row in preview.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(preview.columns):
            value = row[col]
            if pd.isna(value):
                text = ""
            elif isinstance(value, float):
                text = f"{value:.4g}"
            else:
                text = str(value)
            cells[idx].text = text[:300]
    if len(df) > max_rows:
        p = doc.add_paragraph()
        p.add_run(f"Full table contains {len(df)} rows and is provided in the accompanying Excel workbook/source CSV.").italic = True


def write_docx(tables: dict[str, tuple[str, pd.DataFrame, str, str]]) -> Path:
    docx = TABLE_OUT / "Supplementary_Tables_final_20260606.docx"
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)

    doc.add_heading("Supplementary Tables", level=1)
    doc.add_paragraph(
        "This document summarizes the supplementary tables supporting the benchmark manuscript. "
        "The complete machine-readable tables are provided in Supplementary_Tables_final_20260606.xlsx and as individual CSV files."
    )
    doc.add_heading("Table Index", level=2)
    index = table_index_rows(tables)
    add_small_table(doc, index, max_rows=len(index))

    for table_id, (title, df, primary_figures, purpose) in tables.items():
        doc.add_page_break()
        doc.add_heading(f"{supplementary_label(table_id)}. {title}", level=2)
        doc.add_paragraph(f"Machine-readable table ID: {table_id}")
        doc.add_paragraph(f"Primary figure(s)/section(s): {primary_figures}")
        doc.add_paragraph(f"Purpose: {purpose}")
        doc.add_paragraph(f"Rows: {len(df)}; columns: {len(df.columns)}.")
        if len(df) <= 30 and len(df.columns) <= 12:
            add_small_table(doc, df, max_rows=30)
        else:
            add_small_table(doc, df, max_rows=8)
    doc.save(docx)
    return docx


def write_qa(tables: dict[str, tuple[str, pd.DataFrame, str, str]], source_manifest: pd.DataFrame) -> None:
    rows = []
    for table_id, (_, df, _, _) in tables.items():
        rows.append(
            {
                "table_id": table_id,
                "rows": len(df),
                "columns": len(df.columns),
                "duplicate_rows": int(df.duplicated().sum()),
                "empty_columns": "; ".join([col for col in df.columns if df[col].isna().all()]),
            }
        )
    pd.DataFrame(rows).to_csv(QA_OUT / "supplementary_table_qc_summary.csv", index=False)

    method_tables = []
    for table_id, (_, df, _, _) in tables.items():
        method_col = "method_id" if "method_id" in df.columns else "parent_method" if "parent_method" in df.columns else None
        if method_col:
            methods = set(normalize_method(df[method_col].dropna()))
            canonical_seen = len(methods.intersection(CANONICAL_METHODS))
            extras = sorted(methods - set(CANONICAL_METHODS))
            method_tables.append(
                {
                    "table_id": table_id,
                    "method_column": method_col,
                    "canonical_methods_seen": canonical_seen,
                    "extra_methods": "; ".join(extras),
                    "status": "ok" if canonical_seen in {0, 26} or table_id in {"S11_TargetedSensitivity"} else "subset_expected_or_check",
                }
            )
    pd.DataFrame(method_tables).to_csv(QA_OUT / "supplementary_table_method_coverage_audit.csv", index=False)
    source_manifest.to_csv(QA_OUT / "source_file_manifest.csv", index=False)


def write_readme(tables: dict[str, tuple[str, pd.DataFrame, str, str]]) -> None:
    lines = [
        "# Final Supplementary Tables Package",
        "",
        "This folder contains the supplementary tables prepared for Communications Biology submission.",
        "",
        "## Contents",
        "",
        "- `tables/Supplementary_Tables_final_20260606.xlsx`: complete machine-readable supplementary tables.",
        "- `tables/Supplementary_Tables_final_20260606.docx`: manuscript-facing summary document with table index and previews.",
        "- `tables/Supplementary_Table_Captions_final_20260606.docx`: ready-to-paste supplementary table captions.",
        "- `tables/Supplementary_Table_Captions_final_20260606.md`: editable caption and Chinese author-note draft.",
        "- `source_tables/`: one CSV per supplementary table.",
        "- `qa/`: table QC, caption/review map, column audit, method coverage, terminology ledger, and source-file manifest.",
        "- `code/`: script used to regenerate the table package.",
        "",
        "## Table List",
        "",
    ]
    for table_id, (title, df, primary, _) in tables.items():
        lines.append(f"- `{table_id}`: {title} ({len(df)} rows; {primary}).")
    lines.extend(
        [
            "",
            "## Notes",
            "",
            "- `S1_MethodCatalog` contains 26 full-benchmark methods plus three result variants and one targeted scVI reference analysis; only rows labelled `Full benchmark` are counted as the manuscript's 26-method benchmark.",
            "- `S2_Dataset100Atlas` is the 100-dataset backbone: 50 real datasets and 50 simulated datasets.",
            "- `S3_RealDatasetDetail` is aligned to the same 50 real datasets; three large datasets retain atlas-level metadata only because detailed snapshot fields were not present in the available source table.",
            "",
            "## Reproduction",
            "",
            "Run from the project root:",
            "",
            "```powershell",
            'python "Publication\\paper\\revision_tables\\final_supplementary_tables_package_20260606\\make_final_supplementary_tables.py"',
            'python "Publication\\paper\\revision_tables\\final_supplementary_tables_package_20260606\\make_supplementary_table_review.py"',
            "```",
        ]
    )
    (OUT / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    source_manifest = make_source_file_manifest()
    tables: dict[str, tuple[str, pd.DataFrame, str, str]] = {
        "S1_MethodCatalog": (
            "Dimensionality-reduction method catalogue and benchmark scope",
            make_method_catalog(),
            "Figure 1; Supplementary Figure S1",
            "Defines the 26 full-benchmark methods and separates variants/targeted reference analyses.",
        ),
        "S2_Dataset100Atlas": (
            "Manuscript 100-dataset atlas",
            make_dataset_100(),
            "Figure 2; Supplementary Figure S2",
            "Lists the 50 real and 50 simulated datasets used in the two-layer benchmark landscape.",
        ),
        "S3_RealDatasetDetail": (
            "Fifty-real-dataset metadata concordance table",
            make_real_detail(),
            "Figure 2; Supplementary Figure S3",
            "Uses the 50 real datasets in the manuscript atlas as the backbone and marks fields that are available only at atlas level.",
        ),
        "S4_SimulatedAtlas": (
            "Simulated-dataset parameter atlas",
            make_sim_detail(),
            "Figure 2; Supplementary Figures S4 and S10",
            "Records simulated perturbation axes and parameter settings.",
        ),
        "S5_MetricInventory": (
            "Benchmark metric inventory and profile-score components",
            make_metric_inventory(),
            "Figures 3-5; Supplementary Figures S5-S9",
            "Defines score components, local/global metrics, clustering metrics, and raw metric direction.",
        ),
        "S6_ProfileScoreMatrix": (
            "Completed per-method profile-score matrix",
            make_score_matrix(),
            "Figure 3; Supplementary Figure S5",
            "Reports the completed 26-method score matrix after VASC/stability completion.",
        ),
        "S7_StructureCoverage": (
            "Structure-preservation metric coverage summary",
            make_structure_summary(),
            "Figure 4; Supplementary Figures S6-S8",
            "Summarizes local and global structure-preservation records by dataset collection and metric.",
        ),
        "S8_ClusteringCoverage": (
            "Clustering-concordance metric coverage summary",
            make_clustering_summary(),
            "Figure 5; Supplementary Figure S9",
            "Summarizes clustering concordance records by dataset collection, algorithm, and metric.",
        ),
        "S9_RobustnessStability": (
            "Simulated robustness and stability score audit",
            make_robustness_summary(),
            "Figure 7; Supplementary Figure S10",
            "Reports per-method stability components across simulated perturbation axes.",
        ),
        "S10_ScalabilityAudit": (
            "Scalability, completion, and implementation audit",
            make_scalability_audit(),
            "Figure 8; Supplementary Figure S15",
            "Combines scale-completion, runtime/memory, and implementation verification records.",
        ),
        "S11_TargetedSensitivity": (
            "Targeted sensitivity experiment manifest",
            make_revision_experiment_manifest(),
            "Figure 6; Supplementary Figures S11-S14",
            "Summarizes scVI reference, latent-dimension, workflow, and input-gene sensitivity analyses.",
        ),
        "S12_SuppFigurePanelMap": (
            "Supplementary figure panel map",
            make_supp_figure_panel_map(),
            "Supplementary Figures S1-S15",
            "Maps each supplementary panel to its role and source layer.",
        ),
        "S13_SuppEvidenceCoverage": (
            "Supplementary evidence-layer coverage map",
            make_supplementary_evidence_coverage(),
            "Supplementary Figures S6-S10",
            "Summarizes the metric-level evidence layers represented in the supplementary atlas.",
        ),
        "S14_SourceFileManifest": (
            "Source-data file manifest",
            source_manifest,
            "All figures and supplementary tables",
            "Indexes source-data files used by the final figure and table packages.",
        ),
        "S15_DataRouteAudit": (
            "Data availability and source-data route audit",
            make_data_route_audit(source_manifest),
            "Data Availability; all source data",
            "Classifies major data/output classes and flags author-confirmation fields before final submission.",
        ),
    }

    write_csvs(tables)
    xlsx = write_excel(tables)
    docx = write_docx(tables)
    write_qa(tables, source_manifest)
    write_readme(tables)
    shutil.copy2(__file__, CODE_OUT / Path(__file__).name)
    (CODE_OUT / "requirements_minimal.txt").write_text("pandas\nopenpyxl\npython-docx\nnumpy\n", encoding="utf-8")
    print(f"Wrote {len(tables)} supplementary tables")
    print(xlsx)
    print(docx)


if __name__ == "__main__":
    main()
