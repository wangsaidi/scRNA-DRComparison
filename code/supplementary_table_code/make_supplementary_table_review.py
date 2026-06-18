from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt


def find_project_root(start: Path) -> Path:
    for parent in [start.resolve(), *start.resolve().parents]:
        if (parent / "Publication" / "paper").exists() and (parent / "metadata").exists():
            return parent
    raise RuntimeError(f"Could not locate project root from {start}")


ROOT = find_project_root(Path(__file__))
PACKAGE = ROOT / "Publication/paper/revision_tables/final_supplementary_tables_package_20260606"
SOURCE_TABLES = PACKAGE / "source_tables"
TABLES_OUT = PACKAGE / "tables"
QA_OUT = PACKAGE / "qa"
CODE_OUT = PACKAGE / "code"
for path in [TABLES_OUT, QA_OUT, CODE_OUT]:
    path.mkdir(parents=True, exist_ok=True)


TABLE_ORDER = [
    "S1_MethodCatalog",
    "S2_Dataset100Atlas",
    "S3_RealDatasetDetail",
    "S4_SimulatedAtlas",
    "S5_MetricInventory",
    "S6_ProfileScoreMatrix",
    "S7_StructureCoverage",
    "S8_ClusteringCoverage",
    "S9_RobustnessStability",
    "S10_ScalabilityAudit",
    "S11_TargetedSensitivity",
    "S12_SuppFigurePanelMap",
    "S13_SuppEvidenceCoverage",
    "S14_SourceFileManifest",
    "S15_DataRouteAudit",
]


TABLE_META = {
    "S1_MethodCatalog": {
        "label": "Supplementary Table S1",
        "title": "Dimensionality-reduction method catalogue and benchmark scope",
        "figure_links": "Figure 1; Supplementary Figure S1",
        "caption": (
            "Catalogue of dimensionality-reduction methods included in the benchmark. "
            "The table lists the 26 full-benchmark methods, their method families, implementation principles, "
            "implementation languages, software/source links and references, and separately marks result variants "
            "and the targeted scVI reference analysis that are not counted as part of the 26-method benchmark."
        ),
        "core_question": "What exactly counts as the 26-method benchmark, and what is outside that count?",
        "coverage_check": "Contains 26 rows labelled Full benchmark, three result variants, and one targeted scVI reference analysis.",
        "author_action": "Confirm reference formatting and software/source URLs before final submission.",
    },
    "S2_Dataset100Atlas": {
        "label": "Supplementary Table S2",
        "title": "Manuscript 100-dataset atlas",
        "figure_links": "Figure 2; Supplementary Figure S2",
        "caption": (
            "Atlas of the 100 datasets used to frame the benchmark landscape. "
            "The table records 50 real datasets and 50 simulated datasets, including dataset labels, source metadata, "
            "cell counts, sequencing or simulation descriptors, and dataset-counting rules used in the manuscript."
        ),
        "core_question": "Where is the full 100-dataset landscape documented?",
        "coverage_check": "Contains 100 rows: 50 real datasets and 50 simulated datasets.",
        "author_action": "Confirm final public accession wording for reused real datasets in the Data Availability statement.",
    },
    "S3_RealDatasetDetail": {
        "label": "Supplementary Table S3",
        "title": "Fifty-real-dataset metadata concordance table",
        "figure_links": "Figure 2; Supplementary Figure S3",
        "caption": (
            "Concordance table for the 50 real datasets in the manuscript atlas. "
            "Each row is anchored to the real-dataset backbone in Supplementary Table S2 and reports available metadata "
            "for source repository, species, tissue or cell type, condition, cell counts, sequencing technology, "
            "detailed gene/sparsity/cell-type fields where available, and metadata-status notes."
        ),
        "core_question": "How do the real-dataset detail rows align with the 50 real datasets claimed in the manuscript?",
        "coverage_check": "Contains 50 rows; 47 have detailed metadata and three large datasets retain atlas-level metadata only.",
        "author_action": "Do not infer missing detailed fields for Macosko, Zheng-68 k, or Zheng-73 k unless source metadata are verified.",
    },
    "S4_SimulatedAtlas": {
        "label": "Supplementary Table S4",
        "title": "Simulated-dataset parameter atlas",
        "figure_links": "Figure 2; Supplementary Figures S4 and S10",
        "caption": (
            "Parameter atlas for the 50 simulated datasets used in the benchmark. "
            "The table reports simulation axes and parameter values for cell number, gene number, cell-type number, dropout, "
            "batch number, batch strength, differential-expression probability, differential-expression strength and outlier settings."
        ),
        "core_question": "Which simulation perturbations underlie the robustness and stability analyses?",
        "coverage_check": "Contains 50 simulated parameter settings across nine perturbation groups.",
        "author_action": "No action required unless final Methods changes simulation terminology.",
    },
    "S5_MetricInventory": {
        "label": "Supplementary Table S5",
        "title": "Benchmark metric inventory and profile-score components",
        "figure_links": "Figures 3-5; Supplementary Figures S5-S9",
        "caption": (
            "Inventory of benchmark metrics and profile-score components. "
            "The table defines score domains, local and global structure-preservation metrics, clustering-concordance metrics, "
            "raw metric direction and score weights used to construct the profile-score summaries."
        ),
        "core_question": "How were raw metrics mapped into the profile-score framework?",
        "coverage_check": "Contains profile-score components, 34 structure-preservation metrics and clustering metrics across k-means, Louvain and spectral clustering.",
        "author_action": "Use this table when explaining score calculation in Methods and supplementary legends.",
    },
    "S6_ProfileScoreMatrix": {
        "label": "Supplementary Table S6",
        "title": "Completed per-method profile-score matrix",
        "figure_links": "Figure 3; Supplementary Figure S5",
        "caption": (
            "Completed profile-score matrix for the 26 full-benchmark methods. "
            "The table reports local and global structure scores, clustering scores, efficiency scores, stability components "
            "and the overall mean profile score after completing the VASC and stability-score audit."
        ),
        "core_question": "What exact method-level values support Figure 3?",
        "coverage_check": "Contains 26 methods in the canonical order and no missing score columns.",
        "author_action": "No action required unless Figure 3 values are recalculated.",
    },
    "S7_StructureCoverage": {
        "label": "Supplementary Table S7",
        "title": "Structure-preservation metric coverage summary",
        "figure_links": "Figure 4; Supplementary Figures S6-S8",
        "caption": (
            "Coverage summary for local and global structure-preservation metrics. "
            "For each source collection and metric, the table reports the number of datasets, number of methods, record counts "
            "and quartile summaries of metric values used to support Figure 4 and the structure-preservation supplementary atlas."
        ),
        "core_question": "How much structure-preservation evidence exists behind the main Figure 4 summaries?",
        "coverage_check": "Contains 204 source-collection-by-metric rows covering six source collections and 34 metrics.",
        "author_action": "Use source_collection labels in manuscript-facing text rather than internal dataset-block labels.",
    },
    "S8_ClusteringCoverage": {
        "label": "Supplementary Table S8",
        "title": "Clustering-concordance metric coverage summary",
        "figure_links": "Figure 5; Supplementary Figure S9",
        "caption": (
            "Coverage summary for clustering-concordance metrics. "
            "The table summarizes adjusted Rand index, normalized mutual information, homogeneity, completeness and silhouette scores "
            "across k-means, Louvain and spectral clustering for each source collection."
        ),
        "core_question": "How broad is the clustering-concordance evidence behind Figure 5?",
        "coverage_check": "Contains 90 source-collection-by-algorithm-by-metric rows.",
        "author_action": "No action required unless Figure 5 source data are updated.",
    },
    "S9_RobustnessStability": {
        "label": "Supplementary Table S9",
        "title": "Simulated robustness and stability score audit",
        "figure_links": "Figure 7; Supplementary Figure S10",
        "caption": (
            "Per-method robustness and stability audit across simulated perturbation axes. "
            "The table reports method-level scores for perturbations in cell number, gene number, cell-type number, dropout, batch number, "
            "batch strength, differential-expression probability, differential-expression strength and outlier settings, together with the stability median."
        ),
        "core_question": "Which simulated perturbation axes drive the stability conclusions?",
        "coverage_check": "Contains 26 canonical methods and all simulated stability components.",
        "author_action": "No action required unless robustness scoring is recalculated.",
    },
    "S10_ScalabilityAudit": {
        "label": "Supplementary Table S10",
        "title": "Scalability, completion and implementation audit",
        "figure_links": "Figure 8; Supplementary Figure S15",
        "caption": (
            "Audit of scalability, completion status and implementation verification for the 26 full-benchmark methods. "
            "The table records completion across cell-number scales, missing cell levels, largest completed scale, software environment, "
            "installation status, runtime summaries and peak-memory summaries."
        ),
        "core_question": "Which methods completed large-scale runs, and under what implementation conditions?",
        "coverage_check": "Contains 26 canonical methods with scale-completion and efficiency summaries.",
        "author_action": "Confirm final code/software availability wording before submission.",
    },
    "S11_TargetedSensitivity": {
        "label": "Supplementary Table S11",
        "title": "Targeted sensitivity experiment manifest",
        "figure_links": "Figure 6; Supplementary Figures S11-S14",
        "caption": (
            "Manifest and QA summary for targeted sensitivity experiments. "
            "The table documents the scVI reference analysis, latent-dimension sensitivity, visualization-workflow comparison and input-gene sensitivity analyses, "
            "including source-record counts, method scopes, dataset scopes, parameter scopes and Figure 6 QA fields."
        ),
        "core_question": "Which additional sensitivity experiments were run, and how do they relate to scVI and sensitivity controls?",
        "coverage_check": "Contains four experiment-manifest rows and 29 Figure 6 QA-summary rows.",
        "author_action": "Use this table to explain that scVI is a targeted reference analysis rather than part of the prespecified 26-method benchmark.",
    },
    "S12_SuppFigurePanelMap": {
        "label": "Supplementary Table S12",
        "title": "Supplementary figure panel map",
        "figure_links": "Supplementary Figures S1-S15",
        "caption": (
            "Panel-level map of the Supplementary Figure atlas. "
            "The table assigns each panel in Supplementary Figures S1-S15 to its analytical role and source-data layer, allowing readers to trace each supplementary panel to its supporting evidence."
        ),
        "core_question": "What does each supplementary panel contribute?",
        "coverage_check": "Contains 145 panel-level rows covering all 15 supplementary figures.",
        "author_action": "No action required unless supplementary figures are renumbered.",
    },
    "S13_SuppEvidenceCoverage": {
        "label": "Supplementary Table S13",
        "title": "Supplementary evidence-layer coverage map",
        "figure_links": "Supplementary Figures S6-S10",
        "caption": (
            "Coverage map for the metric-level evidence layers represented in the Supplementary Figure atlas. "
            "The table records the supplementary figure, evidence layer, coverage scope, number of evidence units and reader-use note for each major structure, clustering and perturbation analysis layer."
        ),
        "core_question": "Which supplementary figures summarize each major evidence layer?",
        "coverage_check": "Covers local structure, local-manifold, global geometry, clustering and simulated-perturbation evidence layers.",
        "author_action": "No action required unless supplementary figures are renumbered.",
    },
    "S14_SourceFileManifest": {
        "label": "Supplementary Table S14",
        "title": "Source-data file manifest",
        "figure_links": "All figures and supplementary tables",
        "caption": (
            "Manifest of source-data files used to generate the final figures and supplementary tables. "
            "For each file, the table records its source role, relative package path, row count, column count and representative columns, providing an auditable route from source data to displayed results."
        ),
        "core_question": "Where are the source files that support the final figure and table package?",
        "coverage_check": "Indexes all detected source CSV files in the figure/table packages.",
        "author_action": "Confirm final repository or archive location for these files before submission.",
    },
    "S15_DataRouteAudit": {
        "label": "Supplementary Table S15",
        "title": "Data availability and source-data route audit",
        "figure_links": "Data Availability; all source data",
        "caption": (
            "Audit of data and output classes, access routes, package locations and remaining availability actions. "
            "The table separates method metadata, dataset metadata, benchmark outputs, targeted sensitivity experiments, reproduction code and source-file manifests to support the final Data Availability and Code Availability statements."
        ),
        "core_question": "What still needs to be confirmed for data/code availability?",
        "coverage_check": "Contains six data/output classes and identifies repository/accession wording that still requires author confirmation.",
        "author_action": "Confirm final public repository, accession and code-archive DOI wording before submission.",
    },
}


TERMINOLOGY = [
    {
        "canonical_term": "26 full-benchmark methods",
        "definition": "The prespecified benchmark method set counted in the main manuscript.",
        "usage_decision": "Use for the core benchmark only; do not include scVI or result variants in this count.",
    },
    {
        "canonical_term": "targeted scVI reference analysis",
        "definition": "Additional sensitivity analysis used to contextualize scVI without redefining the 26-method benchmark.",
        "usage_decision": "Use in captions and Methods notes when explaining scVI.",
    },
    {
        "canonical_term": "100-dataset atlas",
        "definition": "The manuscript backbone containing 50 real datasets and 50 simulated datasets.",
        "usage_decision": "Use instead of informal phrases such as full landscape or all data unless the count is repeated.",
    },
    {
        "canonical_term": "source collection",
        "definition": "Formal table label for groups of source datasets summarized in structure/clustering coverage tables.",
        "usage_decision": "Use instead of internal labels such as dataset block.",
    },
    {
        "canonical_term": "profile score",
        "definition": "Direction-aligned, normalized and aggregated benchmark score used in Figure 3.",
        "usage_decision": "Use consistently with Supplementary Table S5 and S6.",
    },
]


def read_table(table_id: str) -> pd.DataFrame:
    return pd.read_csv(SOURCE_TABLES / f"{table_id}.csv", keep_default_na=False)


def make_caption_map() -> pd.DataFrame:
    rows = []
    for table_id in TABLE_ORDER:
        df = read_table(table_id)
        meta = TABLE_META[table_id]
        rows.append(
            {
                "supplementary_table": meta["label"],
                "table_id": table_id,
                "title": meta["title"],
                "figure_links": meta["figure_links"],
                "rows": len(df),
                "columns": len(df.columns),
                "caption": meta["caption"],
                "core_question_answered": meta["core_question"],
                "coverage_check": meta["coverage_check"],
                "author_action": meta["author_action"],
            }
        )
    return pd.DataFrame(rows)


def make_column_audit() -> pd.DataFrame:
    rows = []
    for table_id in TABLE_ORDER:
        df = read_table(table_id)
        for column in df.columns:
            values = df[column].astype(str)
            empty = values.str.len().eq(0)
            examples = "; ".join(values[~empty].drop_duplicates().head(5))
            rows.append(
                {
                    "table_id": table_id,
                    "column": column,
                    "non_empty_cells": int((~empty).sum()),
                    "empty_cells": int(empty.sum()),
                    "empty_fraction": round(float(empty.mean()), 4),
                    "example_values": examples,
                }
            )
    return pd.DataFrame(rows)


def write_markdown(caption_map: pd.DataFrame) -> Path:
    md = TABLES_OUT / "Supplementary_Table_Captions_final_20260606.md"
    lines = [
        "# Supplementary Table Captions",
        "",
        "One-sentence argument: The supplementary tables make the benchmark auditable by linking method scope, the 100-dataset atlas, metric construction, figure source data, targeted sensitivity experiments and data/code availability routes to explicit source evidence.",
        "",
        "## Ready-To-Paste Captions",
        "",
    ]
    for _, row in caption_map.iterrows():
        lines.extend(
            [
                f"### {row['supplementary_table']}. {row['title']}",
                "",
                row["caption"],
                "",
                f"Linked evidence: {row['figure_links']}.",
                "",
            ]
        )
    lines.extend(
        [
            "## Chinese Author Notes",
            "",
            "- S1 defines the 26 full-benchmark methods and separately marks scVI as a targeted reference analysis, preventing ambiguity about the 26-method benchmark scope.",
            "- S2/S3/S4 jointly support the 100-dataset atlas, with S3 realigned to the 50 real datasets.",
            "- S5/S6 provide the tabular evidence for Figure 3 score calculation and the final profile-score matrix.",
            "- S7/S8/S9/S10 correspond to structure preservation, clustering concordance, simulated stability and computational efficiency, supporting main Figures 4/5/7/8.",
            "- S11 supports the targeted sensitivity experiments and scVI reference analysis; S12/S13 document panel-level and evidence-layer coverage for the Supplementary Figure atlas; S14/S15 support source data and Data Availability.",
        ]
    )
    md.write_text("\n".join(lines), encoding="utf-8")
    return md


def write_docx(caption_map: pd.DataFrame) -> Path:
    docx = TABLES_OUT / "Supplementary_Table_Captions_final_20260606.docx"
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_heading("Supplementary Table Captions", level=1)
    doc.add_paragraph(
        "The supplementary tables make the benchmark auditable by linking method scope, "
        "the 100-dataset atlas, metric construction, figure source data, targeted sensitivity experiments "
        "and data/code availability routes to explicit source evidence."
    )
    for _, row in caption_map.iterrows():
        doc.add_heading(f"{row['supplementary_table']}. {row['title']}", level=2)
        doc.add_paragraph(row["caption"])
        doc.add_paragraph(f"Linked evidence: {row['figure_links']}.")
    doc.save(docx)
    return docx


def main() -> None:
    caption_map = make_caption_map()
    column_audit = make_column_audit()
    terminology = pd.DataFrame(TERMINOLOGY)

    caption_map.to_csv(QA_OUT / "supplementary_table_caption_and_review_map.csv", index=False, encoding="utf-8-sig")
    column_audit.to_csv(QA_OUT / "supplementary_table_column_audit.csv", index=False, encoding="utf-8-sig")
    terminology.to_csv(QA_OUT / "supplementary_table_terminology_ledger.csv", index=False, encoding="utf-8-sig")
    md = write_markdown(caption_map)
    docx = write_docx(caption_map)
    (CODE_OUT / "make_supplementary_table_review.py").write_text(Path(__file__).read_text(encoding="utf-8"), encoding="utf-8")
    print(f"Wrote caption map for {len(caption_map)} supplementary tables")
    print(md)
    print(docx)


if __name__ == "__main__":
    main()
