from __future__ import annotations

import csv
import re
import time
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path.cwd()
PACKAGE = ROOT / "Publication" / "paper" / "submission_package_communications_biology_20260617"
FINAL = ROOT / "Publication" / "paper" / "final_materials_package_20260608_clean_figlegends"
WORD_DIR = PACKAGE / "02_manuscript" / "word"
QA_DIR = PACKAGE / "08_quality_control"

DOCX_FILES = [
    WORD_DIR / "manuscript_final_clean.docx",
    WORD_DIR / "manuscript_final_red.docx",
]

BODY_STYLE_NAMES = {
    "Normal",
    "Body Text",
    "First Paragraph",
    "Abstract",
    "Image Caption",
    "Caption",
    "Bibliography",
    "Compact",
}

TIMES = "Times New Roman"
BODY_HALF_POINTS = "24"  # 12 pt


def paragraph_text_nodes(paragraph):
    return list(paragraph._p.iter(qn("w:t")))


def clean_marker_artifacts(paragraph) -> int:
    """Remove leftover marker fragments from earlier redline DOCX conversion."""
    changed = 0
    for node in paragraph_text_nodes(paragraph):
        text = node.text or ""
        new = text.replace("[[[REDSTART]]]", "").replace("[[[REDEND]]]", "")
        new = re.sub(r"^\]\]+", "", new)
        if new != text:
            node.text = new
            changed += 1
    return changed


def restore_sentence_spacing(paragraph) -> int:
    """Restore spaces when redline run boundaries collapsed sentence spacing."""
    changed = 0
    previous_last = ""
    for node in paragraph_text_nodes(paragraph):
        text = node.text or ""
        fixed = re.sub(r"([.!?])(?=[A-Z])", r"\1 ", text)
        if fixed != text:
            node.text = fixed
            text = fixed
            changed += 1
        if text:
            first = text[0]
            if (
                previous_last in ".?!"
                and first.isalpha()
                and first.isupper()
                and not text.startswith(" ")
            ):
                node.text = " " + text
                text = node.text
                changed += 1
            previous_last = text.rstrip()[-1] if text.rstrip() else previous_last
    return changed


def ensure_run_rpr(run_element):
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
    return rpr


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run_element, size_12: bool) -> None:
    rpr = ensure_run_rpr(run_element)
    fonts = ensure_child(rpr, "w:rFonts")
    for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
        fonts.set(qn(attr), TIMES)
    if size_12:
        sz = ensure_child(rpr, "w:sz")
        sz.set(qn("w:val"), BODY_HALF_POINTS)
        szcs = ensure_child(rpr, "w:szCs")
        szcs.set(qn("w:val"), BODY_HALF_POINTS)


def set_paragraph_runs_font(paragraph, size_12: bool) -> int:
    count = 0
    for run in paragraph._p.iter(qn("w:r")):
        set_run_font(run, size_12)
        count += 1
    return count


def paragraph_style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def find_reference_start(doc: Document) -> int | None:
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if idx > 100 and (
            re.match(r"^\[?1[\].]?\s+", txt)
            or (txt.startswith("Aibar,") and "SCENIC" in txt)
        ):
            return idx
    return None


def prefix_figure_captions(doc: Document) -> int:
    changed = 0
    fig_no = 1
    for para in doc.paragraphs:
        if paragraph_style_name(para) not in {"Image Caption", "Caption"}:
            continue
        txt = para.text.strip()
        if not txt:
            continue
        if txt.startswith(f"Figure {fig_no}."):
            fig_no += 1
            continue
        if fig_no <= 9 and not txt.startswith("Figure "):
            first_node = next(iter(paragraph_text_nodes(para)), None)
            if first_node is not None:
                first_node.text = f"Figure {fig_no}. " + (first_node.text or "")
                changed += 1
                fig_no += 1
    return changed


def set_style_defaults(doc: Document) -> None:
    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        try:
            style.font.name = TIMES
            style._element.rPr.rFonts.set(qn("w:ascii"), TIMES)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), TIMES)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), TIMES)
            style._element.rPr.rFonts.set(qn("w:cs"), TIMES)
            if style.name in BODY_STYLE_NAMES:
                style.font.size = Pt(12)
        except Exception:
            continue


def ensure_section_review_format(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        sect_pr = section._sectPr
        ln = sect_pr.find(qn("w:lnNumType"))
        if ln is None:
            ln = OxmlElement("w:lnNumType")
            sect_pr.append(ln)
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:start"), "1")
        ln.set(qn("w:restart"), "continuous")
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), "1")


def clear_footer(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        clear_footer(para)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        r = run._r
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        text_run = OxmlElement("w:t")
        text_run.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(text_run)
        r.append(fld_end)
        set_run_font(r, size_12=True)


def set_paragraph_review_spacing(paragraph, is_body: bool) -> None:
    fmt = paragraph.paragraph_format
    if is_body:
        fmt.line_spacing = 2.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)
    elif paragraph_style_name(paragraph).startswith("Heading"):
        fmt.line_spacing = 1.15
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)


def format_docx(docx_path: Path) -> dict[str, int | str]:
    doc = Document(docx_path)
    set_style_defaults(doc)
    ensure_section_review_format(doc)
    add_page_number_footer(doc)
    ref_start = find_reference_start(doc)
    marker_fixes = 0
    spacing_fixes = 0
    caption_prefixes = prefix_figure_captions(doc)
    justified = 0
    double_spaced = 0
    run_fonts = 0

    for idx, para in enumerate(doc.paragraphs):
        style = paragraph_style_name(para)
        is_body = style in BODY_STYLE_NAMES or (ref_start is not None and idx >= ref_start)
        marker_fixes += clean_marker_artifacts(para)
        spacing_fixes += restore_sentence_spacing(para)
        if is_body:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            justified += 1
            double_spaced += 1
        set_paragraph_review_spacing(para, is_body=is_body)
        run_fonts += set_paragraph_runs_font(para, size_12=is_body)

    doc.save(docx_path)
    media_count, media_uncompressed, media_compressed = count_docx_media(docx_path)
    return {
        "file": str(docx_path.relative_to(ROOT)),
        "paragraphs": len(doc.paragraphs),
        "inline_shapes": len(doc.inline_shapes),
        "reference_start": ref_start if ref_start is not None else -1,
        "justified_paragraphs": justified,
        "double_spaced_paragraphs": double_spaced,
        "run_fonts_touched": run_fonts,
        "marker_fixes": marker_fixes,
        "spacing_fixes": spacing_fixes,
        "caption_prefixes": caption_prefixes,
        "media_count": media_count,
        "media_uncompressed": media_uncompressed,
        "media_compressed": media_compressed,
    }


def count_docx_media(docx_path: Path) -> tuple[int, int, int]:
    with zipfile.ZipFile(docx_path) as z:
        media = [i for i in z.infolist() if i.filename.startswith("word/media/") and not i.is_dir()]
        return len(media), sum(i.file_size for i in media), sum(i.compress_size for i in media)


def refresh_manifest(base: Path, out: Path) -> int:
    rows = []
    for path in sorted(base.rglob("*")):
        if path.is_file():
            st = path.stat()
            rows.append(
                {
                    "path": path.relative_to(base).as_posix(),
                    "bytes": st.st_size,
                    "mtime": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(st.st_mtime)),
                }
            )
    out.parent.mkdir(parents=True, exist_ok=True)
    with out.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["path", "bytes", "mtime"])
        writer.writeheader()
        writer.writerows(rows)
    return len(rows)


def main() -> None:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    reports = [format_docx(path) for path in DOCX_FILES]

    # Keep final-materials flat Word files synchronized with the submission package.
    final_docs = FINAL / "submission_documents"
    final_docs.mkdir(parents=True, exist_ok=True)
    for path in DOCX_FILES:
        target = final_docs / path.name
        target.write_bytes(path.read_bytes())

    lines = [
        "# Word Formatting Audit",
        "",
        "Date: 2026-06-17",
        "",
        "The clean and redline Word manuscripts were formatted for Communications Biology / Nature Portfolio revision review: Times New Roman, 12 pt body text, justified body/caption/reference paragraphs, double spacing for review readability, 1-inch margins, continuous line numbering and footer page numbers. Existing images, redline color and numbered citation hyperlinks were preserved where present.",
        "",
        "| File | Paragraphs | Inline figures | Reference start | Justified paragraphs | Double-spaced paragraphs | Runs touched | Marker fixes | Spacing fixes | Caption prefixes | Media count | Media uncompressed bytes |",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for r in reports:
        lines.append(
            f"| {r['file']} | {r['paragraphs']} | {r['inline_shapes']} | {r['reference_start']} | "
            f"{r['justified_paragraphs']} | {r['double_spaced_paragraphs']} | {r['run_fonts_touched']} | {r['marker_fixes']} | "
            f"{r['spacing_fixes']} | {r['caption_prefixes']} | {r['media_count']} | {r['media_uncompressed']} |"
        )
    report_path = QA_DIR / "word_formatting_audit_20260617.md"
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    # Copy audit and script into release/docs/code.
    for target in [
        PACKAGE / "07_source_data_and_code_availability" / "github_release" / "docs" / report_path.name,
        FINAL / "qa" / report_path.name,
        FINAL / "github_release" / "docs" / report_path.name,
    ]:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(report_path.read_text(encoding="utf-8"), encoding="utf-8")

    script_src = ROOT / "Publication" / "paper" / "manuscript_revision_scripts" / "format_word_manuscripts_20260617.py"
    for target in [
        PACKAGE / "07_source_data_and_code_availability" / "github_release" / "code" / "manuscript_code" / script_src.name,
        FINAL / "github_release" / "code" / "manuscript_code" / script_src.name,
    ]:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(script_src.read_text(encoding="utf-8"), encoding="utf-8")

    n1 = refresh_manifest(PACKAGE, PACKAGE / "00_README" / "SUBMISSION_FILE_MANIFEST_20260617.csv")
    (PACKAGE / "00_README" / "SUBMISSION_FILE_MANIFEST.csv").write_text(
        (PACKAGE / "00_README" / "SUBMISSION_FILE_MANIFEST_20260617.csv").read_text(encoding="utf-8"),
        encoding="utf-8",
    )
    n2 = refresh_manifest(
        PACKAGE / "07_source_data_and_code_availability" / "github_release",
        PACKAGE / "07_source_data_and_code_availability" / "github_release" / "FILE_MANIFEST.csv",
    )
    n3 = refresh_manifest(FINAL / "github_release", FINAL / "github_release" / "FILE_MANIFEST.csv")
    print({"reports": reports, "manifest_counts": [n1, n2, n3]})


if __name__ == "__main__":
    main()
