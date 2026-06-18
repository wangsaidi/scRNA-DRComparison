from __future__ import annotations

import re
import shutil
import subprocess
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path.cwd()
OLD_TEX = (
    ROOT
    / "Publication"
    / "paper"
    / "previous_version_reference_20260617"
    / "elsarticle_from_manuscript_final"
    / "elsarticle-template-harv.tex"
)
PACKAGE = ROOT / "Publication" / "paper" / "submission_package_communications_biology_20260617"
CLEAN_DIR = PACKAGE / "02_manuscript" / "latex_clean"
RED_DIR = PACKAGE / "02_manuscript" / "latex_red"
WORD_DIR = PACKAGE / "02_manuscript" / "word"
OUT_DIR = ROOT / "Publication" / "paper" / "manuscript_revision_true_redline_20260617"
QA_DIR = OUT_DIR / "qa"

PANDOC = Path(r"C:\Users\tjwan\AppData\Local\Pandoc\pandoc.exe")
PDFLATEX = Path(r"C:\Users\tjwan\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe")
BIBTEX = Path(r"C:\Users\tjwan\AppData\Local\Programs\MiKTeX\miktex\bin\x64\bibtex.exe")

FIGURES = {
    1: "Figure_1_conceptual_overview.png",
    2: "Figure_2_workflow.png",
    3: "Figure_3_profile_score_audit.png",
    4: "Figure_4_structure_preservation.png",
    5: "Figure_5_clustering_concordance.png",
    6: "Figure_6_targeted_sensitivity_controls.png",
    7: "Figure_7_simulated_robustness.png",
    8: "Figure_8_scalability_reproducibility.png",
    9: "Figure_9_practical_selection.png",
}


def run(cmd: list[str], cwd: Path, log_name: str) -> None:
    proc = subprocess.run(cmd, cwd=cwd, text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    (QA_DIR / log_name).write_text(proc.stdout, encoding="utf-8", errors="ignore")
    if proc.returncode != 0:
        raise RuntimeError(f"Command failed: {' '.join(cmd)}; see {QA_DIR / log_name}")


def strip_latex_for_matching(text: str) -> str:
    text = re.sub(r"%.*", " ", text)
    text = re.sub(r"\\(citep|citet|cite|ref|label|url|href)\{[^{}]*\}", " ", text)
    text = re.sub(r"\\[a-zA-Z*]+(?:\[[^\]]*\])?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"[^A-Za-z0-9.,;:!?() -]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.lower().strip()


def sentence_units(text: str) -> list[str]:
    clean = strip_latex_for_matching(text)
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9a-z])", clean)
    return [p.strip() for p in parts if len(p.split()) >= 5]


def build_old_index(old_tex: str) -> tuple[str, dict[str, list[str]]]:
    old_text = strip_latex_for_matching(old_tex)
    by_anchor: dict[str, list[str]] = {}
    for sent in sentence_units(old_tex):
        words = [w for w in re.findall(r"[a-z0-9-]+", sent) if len(w) > 3]
        for anchor in set(words[:4] + words[-4:]):
            by_anchor.setdefault(anchor, []).append(sent)
    return old_text, by_anchor


def is_old_or_near_old(unit: str, old_text: str, old_index: dict[str, list[str]]) -> bool:
    norm = strip_latex_for_matching(unit)
    all_words = re.findall(r"[a-z0-9-]+", norm)
    words = [w for w in all_words if len(w) > 3]
    if len(all_words) < 3:
        return True
    if norm in old_text:
        return True
    candidates: list[str] = []
    anchors = words if words else all_words
    for anchor in set(anchors[:4] + anchors[-4:]):
        candidates.extend(old_index.get(anchor, []))
    if not candidates:
        return False
    best = max(SequenceMatcher(None, norm, cand).ratio() for cand in candidates[:250])
    threshold = 0.88 if len(words) >= 5 else 0.92
    return best >= threshold


def split_keep_sentence_endings(text: str) -> list[str]:
    if len(strip_latex_for_matching(text).split()) < 5:
        return [text]
    pieces = re.split(r"(?<=[.!?])(\s+)", text)
    out: list[str] = []
    i = 0
    while i < len(pieces):
        if i + 1 < len(pieces):
            out.append(pieces[i] + pieces[i + 1])
            i += 2
        else:
            out.append(pieces[i])
            i += 1
    return [x for x in out if x]


def protect_textcolor_payload(text: str) -> str:
    return text.strip()


def mark_text_segment(text: str, old_text: str, old_index: dict[str, list[str]], stats: dict[str, int]) -> str:
    result: list[str] = []
    for unit in split_keep_sentence_endings(text):
        if is_old_or_near_old(unit, old_text, old_index):
            result.append(unit)
            if len(strip_latex_for_matching(unit).split()) >= 5:
                stats["unchanged_units"] += 1
        else:
            payload = protect_textcolor_payload(unit)
            result.append(r"\textcolor{red}{" + payload + "}")
            if len(strip_latex_for_matching(unit).split()) >= 5:
                stats["red_units"] += 1
    return "".join(result)


def mark_section_command(line: str, old_text: str, old_index: dict[str, list[str]], stats: dict[str, int]) -> str:
    pattern = re.compile(r"^(\\(?:sub)*section\*?\{)(.*)(\})$")
    match = pattern.match(line.strip())
    if not match:
        return line
    title = match.group(2)
    norm_title = strip_latex_for_matching(title)
    title_words = re.findall(r"[a-z0-9-]+", norm_title)
    if norm_title in old_text:
        return line
    if len(title_words) >= 3 and is_old_or_near_old(title, old_text, old_index):
        return line
    stats["red_units"] += 1
    return match.group(1) + r"\textcolor{red}{" + title + "}" + match.group(3)


def mark_caption_line(line: str, old_text: str, old_index: dict[str, list[str]], stats: dict[str, int]) -> str:
    pattern = re.compile(r"^(\\noindent\{\\scriptsize\\textbf\{Figure~\\thefigure\.\}\s*)(.*)(\}\\par)$")
    match = pattern.match(line.strip())
    if not match:
        return line
    return match.group(1) + mark_text_segment(match.group(2), old_text, old_index, stats) + match.group(3)


def make_true_redline(clean_tex: str, old_text: str, old_index: dict[str, list[str]]) -> tuple[str, dict[str, int]]:
    stats = {"red_units": 0, "unchanged_units": 0, "skipped_lines": 0}
    lines = clean_tex.splitlines()
    out: list[str] = []
    in_math = False
    in_bib = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith(r"\bibliographystyle"):
            in_bib = True
        if stripped.startswith(r"\begin{equation") or stripped.startswith(r"\["):
            in_math = True
        if in_bib or in_math or not stripped or stripped.startswith("%"):
            out.append(line)
            stats["skipped_lines"] += 1
        elif re.match(r"\\(?:sub)*section\*?\{", stripped):
            out.append(mark_section_command(line, old_text, old_index, stats))
        elif stripped.startswith(r"\noindent{\scriptsize\textbf{Figure~\thefigure.}"):
            out.append(mark_caption_line(line, old_text, old_index, stats))
        elif stripped.startswith("\\") and not stripped.startswith(r"\noindent{\scriptsize"):
            out.append(line)
            stats["skipped_lines"] += 1
        else:
            out.append(mark_text_segment(line, old_text, old_index, stats))
        if stripped.endswith(r"\end{equation}") or stripped.endswith(r"\]"):
            in_math = False
    red_tex = "\n".join(out) + "\n"
    marker = "\n\\linenumbers\n\n"
    note = (
        "\\noindent\\textcolor{red}{Text in red marks sentences or headings that are new or substantially "
        "modified relative to the original LaTeX submission; unchanged or near-unchanged text is left in black.}\n\n"
    )
    if marker in red_tex and note not in red_tex:
        red_tex = red_tex.replace(marker, marker + note, 1)
    return red_tex, stats


def make_word_friendly_tex(tex: str) -> str:
    for fig_no, filename in FIGURES.items():
        pattern = re.compile(
            r"\\clearpage\s*"
            r"\\begingroup\s*"
            r"\\refstepcounter\{figure\}\s*"
            r"\\label\{fig:fig" + str(fig_no) + r"\}\s*"
            r"\\begin\{center\}\s*"
            r"\\makebox\[\\linewidth\]\[c\]\{\\includegraphics\[[^\]]*\]\{figures/"
            + re.escape(filename)
            + r"\}\}\s*"
            r"\\end\{center\}\s*"
            r"\\vspace\{[-0-9.]+em\}\s*"
            r"\\noindent\{\\scriptsize\\textbf\{Figure~\\thefigure\.\}\s*(?P<caption>.*?)\}\\par\s*"
            r"\\par\\endgroup\s*"
            r"\\clearpage",
            re.S,
        )

        def repl(match: re.Match[str]) -> str:
            caption = match.group("caption").strip()
            return (
                "\\clearpage\n"
                "\\begin{figure}\n"
                "\\centering\n"
                f"\\includegraphics[width=\\textwidth]{{figures/{filename}}}\n"
                f"\\caption{{{caption}}}\n"
                f"\\label{{fig:fig{fig_no}}}\n"
                "\\end{figure}\n"
                "\\clearpage"
            )

        tex, count = pattern.subn(repl, tex, count=1)
        if count != 1:
            raise RuntimeError(f"Could not convert Figure {fig_no} to Word-friendly form.")
    return tex


def replace_textcolor_red_with_markers(tex: str) -> str:
    marker_start = "[[[REDSTART]]]"
    marker_end = "[[[REDEND]]]"
    needle = r"\textcolor{red}{"
    out: list[str] = []
    i = 0
    while i < len(tex):
        j = tex.find(needle, i)
        if j < 0:
            out.append(tex[i:])
            break
        out.append(tex[i:j])
        k = j + len(needle)
        depth = 1
        payload: list[str] = []
        while k < len(tex) and depth:
            ch = tex[k]
            if ch == "\\":
                if k + 1 < len(tex):
                    payload.append(tex[k : k + 2])
                    k += 2
                    continue
            if ch == "{":
                depth += 1
                payload.append(ch)
            elif ch == "}":
                depth -= 1
                if depth:
                    payload.append(ch)
            else:
                payload.append(ch)
            k += 1
        out.append(marker_start + "".join(payload) + marker_end)
        i = k
    return "".join(out)


def compile_latex(tex_name: str) -> None:
    steps = [
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], "true_redline_pdflatex_1.log"),
        ([str(BIBTEX), Path(tex_name).with_suffix(".aux").name], "true_redline_bibtex.log"),
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], "true_redline_pdflatex_2.log"),
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], "true_redline_pdflatex_3.log"),
    ]
    for cmd, log in steps:
        run(cmd, OUT_DIR, log)


def convert_word(tex_name: str, out_docx: Path) -> None:
    cmd = [
        str(PANDOC),
        tex_name,
        "--bibliography=references.bib",
        "--citeproc",
        "--resource-path=.;figures",
        "-o",
        str(out_docx),
    ]
    run(cmd, OUT_DIR, "true_redline_pandoc.log")


def docx_color_count(docx_path: Path) -> tuple[int, int]:
    doc = Document(docx_path)
    red = 0
    total = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                total += 1
                color = run.font.color.rgb
                if color and str(color).upper() == "FF0000":
                    red += 1
    return red, total


def color_docx_marker_ranges(docx_path: Path) -> None:
    start = "[[[REDSTART]]]"
    end = "[[[REDEND]]]"
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = para.text
        if start not in text and end not in text:
            continue
        style = para.style
        pieces: list[tuple[str, bool]] = []
        cursor = 0
        red = False
        while cursor < len(text):
            next_start = text.find(start, cursor)
            next_end = text.find(end, cursor)
            hits = [(idx, kind) for idx, kind in [(next_start, "start"), (next_end, "end")] if idx >= 0]
            if not hits:
                pieces.append((text[cursor:], red))
                break
            idx, kind = min(hits, key=lambda x: x[0])
            if idx > cursor:
                pieces.append((text[cursor:idx], red))
            if kind == "start":
                red = True
                cursor = idx + len(start)
            else:
                red = False
                cursor = idx + len(end)
        para.clear()
        para.style = style
        for piece, is_red in pieces:
            if not piece:
                continue
            run = para.add_run(piece)
            if is_red:
                run.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(docx_path)


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True)
    QA_DIR.mkdir(parents=True)
    shutil.copytree(CLEAN_DIR / "figures", OUT_DIR / "figures")
    for name in ["references.bib", "elsarticle-harv.bst", "elsarticle-num-names.bst", "elsarticle-num.bst"]:
        shutil.copy2(CLEAN_DIR / name, OUT_DIR / name)

    old_tex = OLD_TEX.read_text(encoding="utf-8", errors="ignore")
    clean_tex = (CLEAN_DIR / "manuscript_final_clean.tex").read_text(encoding="utf-8")
    old_text, old_index = build_old_index(old_tex)
    red_tex, stats = make_true_redline(clean_tex, old_text, old_index)
    red_tex_path = OUT_DIR / "manuscript_final_red_true_changes.tex"
    red_tex_path.write_text(red_tex, encoding="utf-8")
    marker_tex = replace_textcolor_red_with_markers(red_tex)
    (OUT_DIR / "manuscript_final_red_true_changes_for_word.tex").write_text(
        make_word_friendly_tex(marker_tex), encoding="utf-8"
    )

    compile_latex(red_tex_path.name)
    convert_word("manuscript_final_red_true_changes_for_word.tex", OUT_DIR / "manuscript_final_red_true_changes.docx")
    color_docx_marker_ranges(OUT_DIR / "manuscript_final_red_true_changes.docx")
    docx_red, docx_total = docx_color_count(OUT_DIR / "manuscript_final_red_true_changes.docx")
    stats["docx_red_runs"] = docx_red
    stats["docx_total_runs"] = docx_total

    report = [
        "# True Redline Generation Report",
        "",
        "Date: 2026-06-17",
        "",
        f"Original LaTeX: `{OLD_TEX.relative_to(ROOT)}`",
        f"Clean revised LaTeX: `{(CLEAN_DIR / 'manuscript_final_clean.tex').relative_to(ROOT)}`",
        "",
        "The previous red version colored the full manuscript body red. This regenerated version marks only sentences or headings that are new or substantially modified relative to the original LaTeX source; unchanged or near-unchanged text remains black.",
        "",
        f"Red/new sentence or heading units: {stats['red_units']}",
        f"Black/near-unchanged sentence units: {stats['unchanged_units']}",
        f"Skipped structural/math/bibliography lines: {stats['skipped_lines']}",
        f"Word red runs / total runs: {docx_red} / {docx_total}",
    ]
    (QA_DIR / "true_redline_generation_report_20260617.md").write_text("\n".join(report) + "\n", encoding="utf-8")
    print(OUT_DIR)


if __name__ == "__main__":
    main()
