from __future__ import annotations

import csv
import re
import time
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path.cwd()
PACKAGE = ROOT / "Publication" / "paper" / "submission_package_communications_biology_20260617"
WORD_DIR = PACKAGE / "02_manuscript" / "word"
LATEX_DIR = PACKAGE / "02_manuscript" / "latex_clean"
QA_DIR = PACKAGE / "08_quality_control"

DOCX_FILES = [
    WORD_DIR / "manuscript_final_clean.docx",
    WORD_DIR / "manuscript_final_red.docx",
]


def parse_bbl_items(bbl_path: Path) -> list[dict[str, str]]:
    text = bbl_path.read_text(encoding="utf-8", errors="ignore")
    pattern = re.compile(r"\\bibitem\[\{(?P<label>.*?)\}\]\{(?P<key>[^}]+)\}", re.S)
    items = []
    for m in pattern.finditer(text):
        raw_label = re.sub(r"\s+", " ", m.group("label").replace("~", " ")).strip()
        year_match = re.search(r"\((\d{4}[a-z]?)\)", raw_label)
        year = year_match.group(1) if year_match else ""
        author_part = raw_label.split("(", 1)[0].strip()
        author_part = author_part.replace(r"\&", "&")
        author_part = re.sub(r"\\[a-zA-Z]+\{([^{}]+)\}", r"\1", author_part)
        author_part = author_part.replace("{", "").replace("}", "")
        author_part = re.sub(r"\s+", " ", author_part).strip()
        key = m.group("key").strip()
        bookmark = "ref_" + re.sub(r"[^A-Za-z0-9_]", "_", key)
        items.append({"key": key, "author": author_part, "year": year, "bookmark": bookmark})
    return items


def citation_variants(item: dict[str, str]) -> list[str]:
    author = item["author"]
    year = item["year"]
    if not author or not year:
        return []
    variants = [
        f"{author} {year}",
        f"{author} ({year})",
        f"{author}, {year}",
    ]
    # Word/citeproc sometimes normalizes "and" to ampersand in parenthetical text.
    if " and " in author:
        amp = author.replace(" and ", " & ")
        variants.extend([f"{amp} {year}", f"{amp} ({year})", f"{amp}, {year}"])
    return sorted(set(variants), key=len, reverse=True)


def find_reference_start(doc: Document, items: list[dict[str, str]]) -> int:
    paragraphs = doc.paragraphs
    first_author = items[0]["author"].split(" ", 1)[0].rstrip(",")
    second_author = items[1]["author"].split(" ", 1)[0].rstrip(",")
    for idx, para in enumerate(paragraphs[:-1]):
        t0 = para.text.strip()
        t1 = paragraphs[idx + 1].text.strip()
        if t0.startswith(first_author + ",") and t1.startswith(second_author + ","):
            return idx
    # Fallback: find the first long reference-like paragraph with a DOI after Supplementary Information.
    for idx, para in enumerate(paragraphs):
        txt = para.text.strip()
        if idx > 100 and txt.startswith(first_author + ",") and "doi.org" in txt:
            return idx
    raise RuntimeError("Could not locate the bibliography start in the Word document.")


def next_bookmark_id(doc: Document) -> int:
    ids = []
    for node in doc.element.body.iter():
        if node.tag == qn("w:bookmarkStart"):
            val = node.get(qn("w:id"))
            if val and val.isdigit():
                ids.append(int(val))
    return (max(ids) + 1) if ids else 1


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    # Avoid duplicate bookmarks when rerun.
    for node in p.iter():
        if node.tag == qn("w:bookmarkStart") and node.get(qn("w:name")) == name:
            return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def build_match_list(text: str, variant_to_anchor: dict[str, str]) -> list[tuple[int, int, str]]:
    matches: list[tuple[int, int, str]] = []
    occupied: list[tuple[int, int]] = []
    for variant in sorted(variant_to_anchor, key=len, reverse=True):
        if len(variant) < 7:
            continue
        pattern = re.compile(r"(?<![A-Za-z])" + re.escape(variant) + r"(?![A-Za-z])")
        for m in pattern.finditer(text):
            s, e = m.span()
            if any(not (e <= os or s >= oe) for os, oe in occupied):
                continue
            matches.append((s, e, variant_to_anchor[variant]))
            occupied.append((s, e))
    return sorted(matches)


def paragraph_has_hyperlink(paragraph) -> bool:
    return any(node.tag == qn("w:hyperlink") for node in paragraph._p.iterchildren())


def make_run(text: str, rpr=None):
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(deepcopy(rpr))
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def make_hyperlink_run(text: str, anchor: str, rpr=None):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    hyperlink.append(make_run(text, rpr))
    return hyperlink


def rebuild_paragraph_with_links(paragraph, matches: list[tuple[int, int, str]]) -> int:
    runs = paragraph.runs
    if not runs or not matches:
        return 0
    segments = []
    cursor = 0
    for run in runs:
        text = run.text
        if not text:
            continue
        start = cursor
        end = cursor + len(text)
        rpr = run._r.rPr
        segments.append((start, end, text, deepcopy(rpr) if rpr is not None else None))
        cursor = end

    p = paragraph._p
    for child in list(p):
        p.remove(child)

    linked = 0
    for start, end, text, rpr in segments:
        pos = start
        while pos < end:
            current = next((m for m in matches if m[0] <= pos < m[1]), None)
            next_boundary = min([m[0] for m in matches if m[0] > pos] + [end])
            if current:
                take_end = min(end, current[1])
                piece = text[pos - start : take_end - start]
                p.append(make_hyperlink_run(piece, current[2], rpr))
                linked += 1
                pos = take_end
            else:
                take_end = min(end, next_boundary)
                piece = text[pos - start : take_end - start]
                if piece:
                    p.append(make_run(piece, rpr))
                pos = take_end
    return linked


def count_docx_media(docx_path: Path) -> tuple[int, int, int]:
    with zipfile.ZipFile(docx_path) as z:
        media = [i for i in z.infolist() if i.filename.startswith("word/media/") and not i.is_dir()]
        return len(media), sum(i.file_size for i in media), sum(i.compress_size for i in media)


def link_docx(docx_path: Path, items: list[dict[str, str]]) -> dict[str, int | str]:
    doc = Document(docx_path)
    ref_start = find_reference_start(doc, items)
    if len(doc.paragraphs) - ref_start < len(items):
        raise RuntimeError(f"Not enough bibliography paragraphs in {docx_path.name}.")

    bookmark_id = next_bookmark_id(doc)
    for offset, item in enumerate(items):
        add_bookmark(doc.paragraphs[ref_start + offset], item["bookmark"], bookmark_id + offset)

    variant_to_anchor: dict[str, str] = {}
    for item in items:
        for variant in citation_variants(item):
            variant_to_anchor[variant] = item["bookmark"]

    linked_paragraphs = 0
    linked_runs = 0
    for idx, para in enumerate(doc.paragraphs[:ref_start]):
        # Do not disturb existing external hyperlinks; DOI/URL links are already useful as-is.
        if paragraph_has_hyperlink(para):
            continue
        text = para.text
        matches = build_match_list(text, variant_to_anchor)
        if not matches:
            continue
        linked = rebuild_paragraph_with_links(para, matches)
        if linked:
            linked_paragraphs += 1
            linked_runs += linked

    doc.save(docx_path)
    media_count, media_uncompressed, media_compressed = count_docx_media(docx_path)
    return {
        "file": str(docx_path.relative_to(ROOT)),
        "reference_start_paragraph": ref_start,
        "bibliography_items": len(items),
        "linked_paragraphs": linked_paragraphs,
        "linked_runs": linked_runs,
        "media_count": media_count,
        "media_uncompressed_bytes": media_uncompressed,
        "media_compressed_bytes": media_compressed,
    }


def refresh_manifest(base: Path) -> int:
    manifest = base / "00_README" / "SUBMISSION_FILE_MANIFEST_20260617.csv"
    rows = []
    for path in sorted(base.rglob("*")):
        if path.is_file():
            st = path.stat()
            rows.append(
                {
                    "path": path.relative_to(base).as_posix(),
                    "bytes": st.st_size,
                    "mtime": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(st.st_mtime)),
                }
            )
    with manifest.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["path", "bytes", "mtime"])
        writer.writeheader()
        writer.writerows(rows)
    shutil_target = base / "00_README" / "SUBMISSION_FILE_MANIFEST.csv"
    shutil_target.write_text(manifest.read_text(encoding="utf-8"), encoding="utf-8")
    return len(rows)


def main() -> None:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    items = parse_bbl_items(LATEX_DIR / "manuscript_final_clean.bbl")
    reports = []
    for docx in DOCX_FILES:
        reports.append(link_docx(docx, items))

    lines = [
        "# Word Citation Hyperlink Audit",
        "",
        "Date: 2026-06-17",
        "",
        "Internal hyperlinks were added from author-year citation text in the manuscript body to bookmarked bibliography paragraphs in the Word documents.",
        "",
        "| File | Bibliography items | Reference start paragraph | Linked paragraphs | Linked runs | Media count | Media uncompressed bytes |",
        "|---|---:|---:|---:|---:|---:|---:|",
    ]
    for report in reports:
        lines.append(
            f"| {report['file']} | {report['bibliography_items']} | {report['reference_start_paragraph']} | "
            f"{report['linked_paragraphs']} | {report['linked_runs']} | {report['media_count']} | {report['media_uncompressed_bytes']} |"
        )
    (QA_DIR / "word_citation_hyperlink_audit_20260617.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    count = refresh_manifest(PACKAGE)
    print({"reports": reports, "manifest_files": count})


if __name__ == "__main__":
    main()
