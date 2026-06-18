from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path.cwd()
MS_DIR = ROOT / "Publication" / "paper" / "manuscript_revision_final_20260609_large_figures_caption_flow"
LATEX_CLEAN = MS_DIR / "latex_clean"
LATEX_RED = MS_DIR / "latex_red"
WORD_OUT = MS_DIR / "word"
QA_OUT = MS_DIR / "qa"

FIG_SRC = (
    ROOT
    / "Publication"
    / "paper"
    / "final_materials_package_20260608_clean_figlegends"
    / "github_release"
    / "figures"
    / "main"
    / "png"
)

PANDOC = Path(r"C:\Users\tjwan\AppData\Local\Pandoc\pandoc.exe")
PDFLATEX = Path(r"C:\Users\tjwan\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe")
BIBTEX = Path(r"C:\Users\tjwan\AppData\Local\Programs\MiKTeX\miktex\bin\x64\bibtex.exe")

FIGURES = {
    1: "Figure_1_conceptual_overview.png",
    2: "Figure_2_workflow.png",
    3: "Figure_3_profile_score_audit.png",
    4: "Figure_4_structure_preservation.png",
    5: "Figure_5_clustering_concordance.png",
    6: "Figure_6_targeted_sensitivity_controls.png",
    7: "Figure_7_simulated_robustness.png",
    8: "Figure_8_scalability_reproducibility.png",
    9: "Figure_9_practical_selection.png",
}


def ensure_dirs() -> None:
    for path in [LATEX_CLEAN, LATEX_RED, WORD_OUT, QA_OUT, LATEX_CLEAN / "figures", LATEX_RED / "figures"]:
        path.mkdir(parents=True, exist_ok=True)


def run(cmd: list[str], cwd: Path, log_name: str) -> int:
    proc = subprocess.run(cmd, cwd=cwd, text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    (QA_OUT / log_name).write_text(proc.stdout, encoding="utf-8", errors="ignore")
    return proc.returncode


def sync_figures() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for fig_no, filename in FIGURES.items():
        src = FIG_SRC / filename
        if not src.exists():
            raise FileNotFoundError(src)
        for target in [LATEX_CLEAN / "figures" / filename, LATEX_RED / "figures" / filename]:
            shutil.copy2(src, target)
        rows.append(
            {
                "figure": f"Figure {fig_no}",
                "file": filename,
                "source": str(src.relative_to(ROOT)),
                "bytes": str(src.stat().st_size),
                "mtime": src.stat().st_mtime_ns.__str__(),
            }
        )
    return rows


def clean_tex_source() -> str:
    tex_path = LATEX_CLEAN / "manuscript_final_clean.tex"
    if not tex_path.exists():
        raise FileNotFoundError(tex_path)
    tex = tex_path.read_text(encoding="utf-8")
    tex = tex.replace(
        "%% Manuscript generated from the user-approved pasted latest text.",
        "%% Manuscript generated from the project-local latest clean LaTeX and synchronized with final figures/tables on 2026-06-16.",
    )
    tex = normalize_figure_reference_style(tex)
    return tex


def normalize_figure_reference_style(text: str) -> str:
    pattern = re.compile(r"(?:Fig\.|Figure)~\\ref\{fig:fig[1-9]\}(?:[a-z](?:--[a-z])?)?")

    def repl(match: re.Match[str]) -> str:
        before = text[: match.start()].rstrip()
        sentence_start = not before or before[-1] in ".?!"
        suffix = match.group(0).split("~", 1)[1]
        return ("Figure~" if sentence_start else "Fig.~") + suffix

    return pattern.sub(repl, text)


def make_red_tex(clean: str) -> str:
    red = clean.replace(r"\title{A systematic benchmark and structured taxonomy of dimensionality-reduction methods for single-cell RNA-seq}", r"\title{{\color{red} A systematic benchmark and structured taxonomy of dimensionality-reduction methods for single-cell RNA-seq}}")
    marker = "\n\\linenumbers\n\n"
    if marker not in red:
        raise RuntimeError("Could not find linenumbers marker in clean manuscript.")
    before, after = red.split(marker, 1)
    body = after
    body = body.replace("\n\\bibliographystyle", "\n}\n\\bibliographystyle", 1)
    return (
        before
        + marker
        + "\\noindent\\textcolor{red}{Text in red marks revised or newly added material relative to the original LaTeX submission.}\n\n"
        + "{\\color{red}\n"
        + body
    )


def compile_latex(tex_dir: Path, tex_name: str, prefix: str) -> None:
    steps = [
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], f"{prefix}_pdflatex_1.log"),
        ([str(BIBTEX), Path(tex_name).with_suffix(".aux").name], f"{prefix}_bibtex.log"),
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], f"{prefix}_pdflatex_2.log"),
        ([str(PDFLATEX), "-interaction=nonstopmode", "-halt-on-error", tex_name], f"{prefix}_pdflatex_3.log"),
    ]
    for cmd, log_name in steps:
        code = run(cmd, tex_dir, log_name)
        if code != 0:
            raise RuntimeError(f"LaTeX build step failed for {tex_name}; see {QA_OUT / log_name}")


def convert_word(tex_dir: Path, tex_name: str, out_docx: Path, prefix: str) -> None:
    cmd = [
        str(PANDOC),
        tex_name,
        "--bibliography=references.bib",
        "--citeproc",
        "--resource-path=.;figures",
        "-o",
        str(out_docx),
    ]
    code = run(cmd, tex_dir, f"{prefix}_pandoc.log")
    if code != 0:
        raise RuntimeError(f"Pandoc failed for {tex_name}; see {QA_OUT / f'{prefix}_pandoc.log'}")


def prefix_word_captions(docx_path: Path) -> None:
    doc = Document(docx_path)
    starts = {
        1: "Conceptual taxonomy",
        2: "Overall workflow",
        3: "Direction-consistent score construction",
        4: "Local-neighborhood and global-geometry preservation",
        5: "Annotation-concordance and clustering-workflow dependence",
        6: "Sensitivity controls",
        7: "Robustness landscape",
        8: "Computational scalability",
        9: "Practical method-selection synthesis",
    }
    for para in doc.paragraphs:
        txt = para.text.strip()
        for idx, start in starts.items():
            if txt.startswith(start) and not txt.startswith(f"Figure {idx}."):
                if para.runs:
                    para.runs[0].text = f"Figure {idx}. " + para.runs[0].text
                break
    doc.save(docx_path)


def color_docx_red(docx_path: Path) -> None:
    doc = Document(docx_path)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run.font.color.rgb = RGBColor(255, 0, 0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(docx_path)


def collect_supplementary_numbers(text: str, kind: str) -> list[int]:
    found: set[int] = set()
    prefix = rf"Supplementary {kind}s?\.?"
    for start, end in re.findall(prefix + r"\s*S([0-9]+)\s*-{1,2}\s*S?([0-9]+)", text):
        found.update(range(int(start), int(end) + 1))
    found.update(int(x) for x in re.findall(prefix + r"\s*S([0-9]+)", text))
    return sorted(found)


def qa_report(figure_rows: list[dict[str, str]]) -> None:
    clean = (LATEX_CLEAN / "manuscript_final_clean.tex").read_text(encoding="utf-8")
    bib = (LATEX_CLEAN / "references.bib").read_text(encoding="utf-8", errors="ignore")
    cite_keys = sorted(
        set(k.strip() for m in re.finditer(r"\\cite[pt]?\{([^}]+)\}", clean) for k in m.group(1).split(","))
    )
    bib_keys = set(re.findall(r"@\w+\{([^,]+),", bib))
    missing = [k for k in cite_keys if k not in bib_keys]
    fig_refs = sorted(set(int(x) for x in re.findall(r"fig:fig([1-9])", clean)))
    supp_figs = collect_supplementary_numbers(clean, "Fig")
    supp_tables = collect_supplementary_numbers(clean, "Table")

    lines = [
        "# Latest Manuscript QA",
        "",
        "Date: 2026-06-16",
        "",
        f"Citation keys: {len(cite_keys)}",
        f"Missing citation keys: {missing or 'none'}",
        f"Main figure labels/references detected: {fig_refs}",
        f"Supplementary figure numbers directly detected: {supp_figs}",
        f"Supplementary table numbers directly detected: {supp_tables}",
        "",
        "Expected main figures: 1-9",
        "Expected supplementary figures: 1-15 (ranges such as S1-S5 may appear in text)",
        "Expected supplementary tables: 1-15 (ranges such as S1-S6 may appear in text)",
        "",
        "## Synchronized Main-Figure Assets",
        "",
        "| Figure | File | Source | Bytes |",
        "|---|---|---|---:|",
    ]
    for row in figure_rows:
        lines.append(f"| {row['figure']} | {row['file']} | {row['source']} | {row['bytes']} |")
    (QA_OUT / "coverage_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    required = [
        not missing,
        fig_refs == list(range(1, 10)),
        supp_figs == list(range(1, 16)),
        supp_tables == list(range(1, 16)),
    ]
    if not all(required):
        raise RuntimeError("Manuscript QA failed; see coverage_report.md")


def main() -> None:
    ensure_dirs()
    figure_rows = sync_figures()

    clean = clean_tex_source()
    (LATEX_CLEAN / "manuscript_final_clean.tex").write_text(clean, encoding="utf-8")
    (LATEX_RED / "references.bib").write_text((LATEX_CLEAN / "references.bib").read_text(encoding="utf-8"), encoding="utf-8")
    (LATEX_RED / "manuscript_final_red.tex").write_text(make_red_tex(clean), encoding="utf-8")

    compile_latex(LATEX_CLEAN, "manuscript_final_clean.tex", "clean_update_20260616")
    compile_latex(LATEX_RED, "manuscript_final_red.tex", "red_update_20260616")

    clean_docx = WORD_OUT / "manuscript_final_clean.docx"
    red_docx = WORD_OUT / "manuscript_final_red.docx"
    convert_word(LATEX_CLEAN, "manuscript_final_clean.tex", clean_docx, "clean_update_20260616")
    convert_word(LATEX_RED, "manuscript_final_red.tex", red_docx, "red_update_20260616")
    prefix_word_captions(clean_docx)
    prefix_word_captions(red_docx)
    color_docx_red(red_docx)
    qa_report(figure_rows)
    print(MS_DIR)


if __name__ == "__main__":
    main()
